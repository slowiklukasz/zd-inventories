import sys
import os
from PyQt5.QtWidgets import *
from PyQt5.QtSql import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from ui_modules.zd_catalog_ui import *
from ui_modules.zd_inventory_ui import *
from ui_modules.zd_contractor_ui import *
from ui_modules.zd_return_ui import *
from ui_modules.zd_mail_ui import *
from ui_modules.zd_tags_return_ui import *
from ui_modules.zd_equipment_ui import *
from ui_modules.zd_print_ui import *
import datetime
from docx import Document
import smtplib


class DlgMain(QDialog, Ui_dlg_catalog):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle("REJESTR WYPOŻYCZEŃ SPRZĘTU")
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)

        self.tbv_finished.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbv_finished.setSelectionMode(QAbstractItemView.SingleSelection)
        self.tbv_progress.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbv_progress.setSelectionMode(QAbstractItemView.SingleSelection)
        self.tbv_history.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbv_history.setSelectionMode(QAbstractItemView.SingleSelection)
        self.tbv_arbotags.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbv_arbotags.setSelectionMode(QAbstractItemView.SingleSelection)
        self.tbv_hammers.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbv_hammers.setSelectionMode(QAbstractItemView.SingleSelection)
        self.tbv_containers.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tbv_containers.setSelectionMode(QAbstractItemView.SingleSelection)

        self.tab_inv.setStyleSheet("QTabBar::tab { height: 150px; width: 25px; }")

        self.led_id.setEnabled(False)
        self.led_name.setEnabled(False)
        self.led_phone.setEnabled(False)
        self.led_email.setEnabled(False)
        self.ted_note.setEnabled(False)
        self.btn_del.setEnabled(False)

        self.btn_ctr_save.setVisible(False)
        self.btn_ctr_cancel.setVisible(False)

        # DATABASE OPEN
        db = QSqlDatabase.addDatabase("QSQLITE")
        db_path = os.getcwd() + r"\{}".format("ZD_CATALOG.db")
        db.setDatabaseName(db_path)

        # POPULATE TABLE VIEWS
        if db.open():
            if "inventories" not in db.tables():
                self.create_inventories_table()
            if "contractors" not in db.tables():
                self.create_contractors_table()
            if "hammers" not in db.tables():
                self.create_hammers_table()
            if "containers" not in db.tables():
                self.create_containers_table()
            if "arbotags" not in db.tables():
                self.create_arbotags_table()
            self.populate_tbv_progress()
            self.populate_tbv_finished()
            self.populate_tbv_canceled()
            self.populate_lsv_contractors()
            self.populate_tbv_history()
            self.populate_tbv_arbotags()
            self.populate_tbv_hammers()
            self.populate_tbv_containers()
            self.lyt_contractors_info_reset()
            self.evt_tab_inv_changed()
        else:
            QMessageBox.critical(self, "Błąd bazy danych", "Brak dostępu do bazy danych")

        # SIGNALS SECTION
        self.btn_inv.clicked.connect(self.evt_btn_inv_clicked)
        self.btn_contr.clicked.connect(self.evt_btn_contr_clicked)
        self.btn_tag.clicked.connect(self.evt_btn_tag_clicked)

        self.btn_add.clicked.connect(self.evt_btn_add_clicked)
        self.btn_mod.clicked.connect(self.evt_btn_mod_clicked)
        self.btn_cancel.clicked.connect(self.evt_btn_cancel_clicked)
        self.btn_del.clicked.connect(self.evt_btn_del_clicked)
        self.btn_tags_return.clicked.connect(self.evt_btn_tags_return)
        self.btn_tags_cancel.clicked.connect(self.evt_btn_tags_cancel)
        self.btn_atags_return.clicked.connect(self.evt_btn_atags_return)
        self.btn_atags_cancel.clicked.connect(self.evt_btn_atags_cancel)
        self.btn_return.clicked.connect(self.evt_btn_return_clicked)
        self.btn_resume.clicked.connect(self.evt_btn_resume_clicked)
        self.btn_add_hammer.clicked.connect(self.evt_btn_add_hammer_clicked)
        self.btn_add_cont.clicked.connect(self.evt_btn_add_cont_clicked)
        self.btn_del_hammer.clicked.connect(self.evt_btn_del_hammer_clicked)
        self.btn_del_cont.clicked.connect(self.evt_btn_del_cont_clicked)

        self.btn_mail.clicked.connect(self.evt_btn_mail_clicked)
        self.btn_print.clicked.connect(self.evt_btn_print_clicked)

        self.tbv_progress.doubleClicked.connect(self.evt_btn_mod_clicked)
        self.tbv_finished.doubleClicked.connect(self.evt_btn_mod_clicked)
        self.lsv_contractors.doubleClicked.connect(self.evt_btn_ctr_mod_clicked)

        self.btn_ctr_add.clicked.connect(self.evt_btn_ctr_add_clicked)
        self.btn_ctr_mod.clicked.connect(self.evt_btn_ctr_mod_clicked)
        self.btn_ctr_del.clicked.connect(self.evt_btn_ctr_del_clicked)
        self.btn_ctr_save.clicked.connect(self.evt_btn_ctr_save_clicked)
        self.btn_ctr_cancel.clicked.connect(self.evt_btn_ctr_cancel_clicked)

        self.lsv_contractors.clicked.connect(self.evt_lsv_contractors_clicked)
        self.tab_inv.currentChanged.connect(self.evt_tab_inv_changed)

    # TABLES CREATION
    @staticmethod
    def create_inventories_table():
        sql = """
            CREATE TABLE IF NOT EXISTS inventories (
                id INTEGER PRIMARY KEY,
                contract_nb TEXT,
                id_contractor INTEGER NOT NULL,
                id_hammer INTEGER,
                id_container INTEGER,
                id_arbotags INTEGER,
                date_from DATE,
                date_to DATE,
                date_return DATE,
                date_hammer_return DATE,
                date_container_return DATE,
                note TEXT,
                status TEXT NOT NULL
            )
        """
        query = QSqlQuery()
        query.exec(sql)
        print(query.lastError().text())

        query.exec("""  INSERT INTO inventories VALUES (1, '2/2022', 5 ,1, 1, 1,
                        '2021-12-01', '2021-12-18', '2021-12-15', NULL, NULL, 'brak uwag', 'ZAKOŃCZONE')""")

        query.exec("""  INSERT INTO inventories VALUES (2, '3/2022',6 ,2, 2, 2,
                        '2021-12-01', '2021-12-18','2021-12-19', NULL, NULL,'brak uwag', 'ZAKOŃCZONE')""")

        query.exec("""  INSERT INTO inventories VALUES (5, '1/2022', 7 ,3, 3, 3,
                       '2021-12-01', '2021-12-20', NULL, NULL, NULL, 'brak uwag', 'W TRAKCIE')""")

    @staticmethod
    def create_contractors_table():
        sql = """
            CREATE TABLE IF NOT EXISTS contractors (
                id INTEGER PRIMARY KEY,
                name TEXT,
                phone TEXT,
                email TEXT,
                note TEXT
            )
        """
        query = QSqlQuery()
        query.exec(sql)

        query.exec("INSERT INTO contractors VALUES (5,'wykonawca5','111-111-111', 'mail1@gmail.com','brak uwag')")
        query.exec("INSERT INTO contractors VALUES (6,'wykonawca6','222-222-222', 'mail2@gmail.com','brak uwag')")
        query.exec("INSERT INTO contractors VALUES (7,'wykonawca7','333-333-333', 'mail3@gmail.com','brak uwag')")

    @staticmethod
    def create_hammers_table():
        sql = """
            CREATE TABLE IF NOT EXISTS hammers (
                id INTEGER PRIMARY KEY,
                number TEXT NOT NULL,
                note TEXT,
                status TEXT NOT NULL
            )"""

        query = QSqlQuery()
        query.exec(sql)
        query.exec("insert into hammers values(1,'-', '', '---')")
        query.exec("insert into hammers values(2,'nr 1', '','ODDANY')")
        query.exec("insert into hammers values(3,'nr 2','', 'WYPOŻYCZONY')")
        query.exec("insert into hammers values(4,'nr 3', '', 'ODDANY')")
        query.exec("insert into hammers values(5,'nr 4', '', 'ODDANY')")
        query.exec("insert into hammers values(6,'nr 5', '', 'ODDANY')")
        query.exec("insert into hammers values(7,'nr 6', '', 'ODDANY')")
        query.exec("insert into hammers values(8,'nr 7', '', 'ODDANY')")
        query.exec("insert into hammers values(9,'nr 8', '', 'ODDANY')")
        query.exec("insert into hammers values(10,'nr 9', '', 'ODDANY')")

    @staticmethod
    def create_containers_table():
        sql = """
            CREATE TABLE IF NOT EXISTS containers (
                id INTEGER PRIMARY KEY,
                number TEXT NOT NULL,
                note TEXT,
                status TEXT NOT NULL
            )"""

        query = QSqlQuery()
        query.exec(sql)
        query.exec("insert into containers values(1,'-', '', '---')")
        query.exec("insert into containers values(2,'nr 1','', 'ODDANY')")
        query.exec("insert into containers values(3,'nr 2', '', 'WYPOŻYCZONY')")
        query.exec("insert into containers values(4,'nr 3', '', 'ODDANY')")
        query.exec("insert into containers values(5,'nr 4', '', 'ODDANY')")
        query.exec("insert into containers values(6,'nr 5', '', 'ODDANY')")

    @staticmethod
    def create_arbotags_table():
        sql = """
            CREATE TABLE IF NOT EXISTS arbotags (
                id INTEGER PRIMARY KEY,
                id_contractor INTEGER,
                arbotags TEXT,
                arbotags_returned TEXT,
                note TEXT
            )"""

        query = QSqlQuery()
        query.exec(sql)
        query.exec("insert into arbotags values(1, 5, '100000-110000\n110005-110014', '110015-110050', '')")
        query.exec("insert into arbotags values(2, 6,'123000-123500\n125001-125604\n125801-125804', '', '')")
        query.exec("insert into arbotags values(3, 7, '126001-126647\n127000-127500', '123000-123200', '')")
        print(query.lastError().text())

    # POPULATING TABLE/LIST VIEW
    def populate_tbv_progress(self):
        self.mdl_progress = ProgressTableModel()
        query = QSqlQuery("""
            SELECT 
                i.id AS ID_INVENTORY, i.contract_nb, c.id AS ID_CONTRACTOR, i.id_contractor, c.name, 
                h.id AS ID_HAMMER, i.id_hammer, h.number AS NB_HAMMER,h.status AS STATUS_HAMMER,
                ct.id AS ID_CONTAINER,i.id_container, ct.number AS NB_CONTAINER, ct.status AS STATUS_CONTAINER,
                i.id_arbotags, a.arbotags AS ARBOTAGS, a.arbotags_returned AS ARBOTAGS_RETURNED,
                i.date_from, i.date_to, i.date_return, 
                date_hammer_return, date_container_return,
                i.note, i.status AS INV_STATUS
            FROM inventories i
            LEFT JOIN contractors c
                ON i.id_contractor = c.id
            LEFT JOIN arbotags a
                ON i.id_arbotags = a.id
            LEFT JOIN hammers h
                ON i.id_hammer = h.id
            LEFT JOIN containers ct
                ON i.id_container = ct.id
            WHERE i.date_return IS NULL AND i.status = "W TRAKCIE"
            ORDER BY i.date_to ASC
             """)
        self.mdl_progress.setQuery(query)
        self.mdl_progress.select()
        self.tbv_progress.setModel(self.mdl_progress)

        self.tbv_progress.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_progress.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.tbv_progress.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        self.mdl_progress.setHeaderData(0, Qt.Orientation.Horizontal, "Id")
        self.mdl_progress.setHeaderData(1, Qt.Orientation.Horizontal, "NR UMOWY")
        self.mdl_progress.setHeaderData(4, Qt.Orientation.Horizontal, "UŻYTKOWNIK")
        self.mdl_progress.setHeaderData(7, Qt.Orientation.Horizontal, "NR\n MŁOTKA")
        self.mdl_progress.setHeaderData(11, Qt.Orientation.Horizontal, "NR\n POJEMNIKA")
        self.mdl_progress.setHeaderData(14, Qt.Orientation.Horizontal, "ARBOTAGI")
        self.mdl_progress.setHeaderData(15, Qt.Orientation.Horizontal, "ARBOTAGI\nZWRÓCONE")
        self.mdl_progress.setHeaderData(16, Qt.Orientation.Horizontal, "DATA OD")
        self.mdl_progress.setHeaderData(17, Qt.Orientation.Horizontal, "DATA DO")
        self.mdl_progress.setHeaderData(18, Qt.Orientation.Horizontal, "DATA ZWROTU")

        self.tbv_progress.setColumnHidden(2, True)
        self.tbv_progress.setColumnHidden(3, True)
        self.tbv_progress.setColumnHidden(5, True)
        self.tbv_progress.setColumnHidden(6, True)
        self.tbv_progress.setColumnHidden(8, True)
        self.tbv_progress.setColumnHidden(9, True)
        self.tbv_progress.setColumnHidden(10, True)
        self.tbv_progress.setColumnHidden(12, True)
        self.tbv_progress.setColumnHidden(13, True)
        self.tbv_progress.setColumnHidden(19, True)
        self.tbv_progress.setColumnHidden(20, True)
        self.tbv_progress.setColumnHidden(21, True)
        self.tbv_progress.setColumnHidden(22, True)

    def populate_tbv_finished(self):
        # self.mdl_finished = QSqlTableModel()
        self.mdl_finished = ModifiedTableModel()
        query = QSqlQuery("""
            SELECT 
                i.id AS ID_INVENTORY, i.contract_nb, c.id AS ID_CONTRACTOR, i.id_contractor, c.name, 
                h.id AS ID_HAMMER, i.id_hammer, h.number AS NB_HAMMER, h.status AS STATUS_HAMMER,
                ct.id AS ID_CONTAINER, i.id_container, ct.number AS NB_CONTAINER, ct.status AS STATUS_CONTAINER,
                i.id_arbotags, a.arbotags AS ARBOTAGS,a.arbotags_returned AS ARBOTAGS_RETURNED,
                i.date_from, i.date_to, i.date_return, 
                i.note, i.status AS INV_STATUS
            FROM inventories i
            LEFT JOIN contractors c
                ON c.id=i.id_contractor
            LEFT JOIN arbotags a
                ON i.id_arbotags = a.id
            LEFT JOIN hammers h
                ON h.id=i.id_hammer
            LEFT JOIN containers ct
                ON ct.id=i.id_container
            WHERE i.date_return IS NOT NULL AND i.status = "ZAKOŃCZONE"
            ORDER BY i.date_return DESC
             """)

        self.mdl_finished.setQuery(query)
        self.mdl_finished.select()
        self.tbv_finished.setModel(self.mdl_finished)

        self.tbv_finished.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_finished.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.tbv_finished.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        self.mdl_finished.setHeaderData(0, Qt.Orientation.Horizontal, "Id")
        self.mdl_finished.setHeaderData(1, Qt.Orientation.Horizontal, "NR UMOWY")
        self.mdl_finished.setHeaderData(4, Qt.Orientation.Horizontal, "UŻYTKOWNIK")
        self.mdl_finished.setHeaderData(7, Qt.Orientation.Horizontal, "NR\n MŁOTKA")
        self.mdl_finished.setHeaderData(11, Qt.Orientation.Horizontal, "NR\n POJEMNIKA")
        self.mdl_finished.setHeaderData(14, Qt.Orientation.Horizontal, "ARBOTAGI")
        self.mdl_finished.setHeaderData(15, Qt.Orientation.Horizontal, "ARBOTAGI\nZWRÓCONE")
        self.mdl_finished.setHeaderData(16, Qt.Orientation.Horizontal, "DATA OD")
        self.mdl_finished.setHeaderData(17, Qt.Orientation.Horizontal, "DATA DO")
        self.mdl_finished.setHeaderData(18, Qt.Orientation.Horizontal, "DATA ZWROTU")

        self.tbv_finished.setColumnHidden(2, True)
        self.tbv_finished.setColumnHidden(3, True)
        self.tbv_finished.setColumnHidden(5, True)
        self.tbv_finished.setColumnHidden(6, True)
        self.tbv_finished.setColumnHidden(8, True)
        self.tbv_finished.setColumnHidden(9, True)
        self.tbv_finished.setColumnHidden(10, True)
        self.tbv_finished.setColumnHidden(12, True)
        self.tbv_finished.setColumnHidden(13, True)
        self.tbv_finished.setColumnHidden(19, True)
        self.tbv_finished.setColumnHidden(20, True)

    def populate_tbv_canceled(self):
        # self.mdl_canceled = QSqlTableModel()
        self.mdl_canceled = ModifiedTableModel()
        query = QSqlQuery("""
            SELECT 
                i.id AS ID_INVENTORY, c.id AS ID_CONTRACTOR, i.id_contractor, c.name, 
                h.id AS ID_HAMMER, i.id_hammer, h.number AS NB_HAMMER, h.status AS STATUS_HAMMER,
                ct.id AS ID_CONTAINER, i.id_container, ct.number AS NB_CONTAINER, ct.status AS STATUS_CONTAINER,
                i.id_arbotags, i.date_from, i.date_to, i.date_return, i.note, i.status AS INV_STATUS
            FROM inventories i
            LEFT JOIN contractors c
                ON c.id=i.id_contractor
            LEFT JOIN arbotags a
                ON i.id_arbotags = a.id
            LEFT JOIN hammers h
                ON h.id=i.id_hammer
            LEFT JOIN containers ct
                ON ct.id=i.id_container
            WHERE i.status = 'ANULOWANE'
            ORDER BY i.date_return DESC
             """)

        self.mdl_canceled.setQuery(query)
        print(query.lastError().text())
        self.mdl_canceled.select()
        self.tbv_canceled.setModel(self.mdl_canceled)

        self.tbv_canceled.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_canceled.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.tbv_canceled.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        self.mdl_canceled.setHeaderData(0, Qt.Orientation.Horizontal, "Id")
        self.mdl_canceled.setHeaderData(3, Qt.Orientation.Horizontal, "UŻYTKOWNIK")
        self.mdl_canceled.setHeaderData(6, Qt.Orientation.Horizontal, "NR\n MŁOTKA")
        self.mdl_canceled.setHeaderData(10, Qt.Orientation.Horizontal, "NR\n POJEMNIKA")
        self.mdl_canceled.setHeaderData(13, Qt.Orientation.Horizontal, "DATA OD")
        self.mdl_canceled.setHeaderData(14, Qt.Orientation.Horizontal, "DATA DO")
        self.mdl_canceled.setHeaderData(15, Qt.Orientation.Horizontal, "DATA ZWROTU")

        self.tbv_canceled.setColumnHidden(1, True)
        self.tbv_canceled.setColumnHidden(2, True)
        self.tbv_canceled.setColumnHidden(4, True)
        self.tbv_canceled.setColumnHidden(5, True)
        self.tbv_canceled.setColumnHidden(7, True)
        self.tbv_canceled.setColumnHidden(8, True)
        self.tbv_canceled.setColumnHidden(9, True)
        self.tbv_canceled.setColumnHidden(11, True)
        self.tbv_canceled.setColumnHidden(12, True)
        self.tbv_canceled.setColumnHidden(16, True)
        self.tbv_canceled.setColumnHidden(17, True)
        self.tbv_canceled.setColumnHidden(18, True)

    def populate_lsv_contractors(self):
        self.mdl_contractors = QSqlQueryModel()
        self.mdl_contractors.setQuery("SELECT name, id, phone, email, note FROM contractors")
        self.lsv_contractors.setModel(self.mdl_contractors)

    def populate_tbv_history(self):
        # self.mdl_history = QSqlTableModel()
        self.mdl_history = ModifiedTableModel()
        query = QSqlQuery("""
            SELECT 
                i.id AS ID, 
                h.number AS 'NR \nMŁOTA',
                ct.number AS 'NR \nPOJEMNIKA',
                a.arbotags AS ARBOTAGS,
                a.arbotags_returned AS ARBOTAGS_RETURNED,
                i.date_from AS 'DATA \nOD', 
                i.date_to AS 'DATA \nDO', 
                i.date_return AS 'DATA \nZWROTU',
                i.status AS STATUS
            FROM inventories i
            LEFT JOIN contractors c
                ON i.id_contractor = c.id
            LEFT JOIN hammers h
                ON i.id_hammer = h.id
            LEFT JOIN containers ct
                ON i.id_container = ct.id
            LEFT JOIN arbotags a
                ON i.id_arbotags = a.id
            WHERE i.status <> 'ANULOWANE' AND c.id = '{}'
            ORDER BY i.date_to DESC
             """.format(self.led_id.text()))
        self.mdl_history.setQuery(query)
        self.mdl_history.select()
        self.tbv_history.setModel(self.mdl_history)

        self.tbv_history.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_history.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.tbv_history.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        self.mdl_history.setHeaderData(3, Qt.Orientation.Horizontal, "ARBOTAGI")
        self.mdl_history.setHeaderData(4, Qt.Orientation.Horizontal, "ARBOTAGI\nZWRÓCONE")

        self.tbv_history.setColumnHidden(0, True)

    def populate_tbv_arbotags(self):
        # self.mdl_arbotags = QSqlQueryModel()
        self.mdl_arbotags = ModifiedSqlQueryModel()
        self.mdl_arbotags.setQuery("""
            SELECT 
                a.id, c.name, a.arbotags, a.arbotags_returned
            FROM arbotags a
            LEFT JOIN contractors c
                ON a.id_contractor = c.id
            WHERE a.arbotags <> '' or  a.arbotags_returned <> ''
            ORDER BY a.arbotags DESC
            """)
        self.tbv_arbotags.setModel(self.mdl_arbotags)

        self.mdl_arbotags.setHeaderData(1, Qt.Orientation.Horizontal, "UŻYTKOWNIK")
        self.mdl_arbotags.setHeaderData(2, Qt.Orientation.Horizontal, "ARBOTAGI WYDANE")
        self.mdl_arbotags.setHeaderData(3, Qt.Orientation.Horizontal, "ARBOTAGI ZWRÓCONE")

        self.tbv_arbotags.setColumnHidden(0, True)

        self.tbv_arbotags.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_arbotags.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

    def populate_tbv_hammers(self):
        # self.mdl_hammers = QSqlQueryModel()
        self.mdl_hammers = ModifiedSqlQueryModel()
        self.mdl_hammers.setQuery("""
            SELECT h.id, h.number, c.name AS LAST_BORROWER, h.status
            FROM hammers h
            LEFT JOIN inventories i
                ON h.id = i.id_hammer
            LEFT JOIN contractors c
                ON i.id_contractor = c.id
            WHERE h.id <> 1
            GROUP BY h.id
            """)
        self.tbv_hammers.setModel(self.mdl_hammers)

        self.mdl_hammers.setHeaderData(1, Qt.Orientation.Horizontal, "NR MŁOTKA")
        self.mdl_hammers.setHeaderData(2, Qt.Orientation.Horizontal, "OSTATNI\nUŻYTKOWNIK")
        self.mdl_hammers.setHeaderData(3, Qt.Orientation.Horizontal, "STATUS")

        self.tbv_hammers.setColumnHidden(0, True)

        self.tbv_hammers.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_hammers.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.tbv_hammers.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

    def populate_tbv_containers(self):
        # self.mdl_containers = QSqlQueryModel()
        self.mdl_containers = ModifiedSqlQueryModel()
        self.mdl_containers.setQuery("""
            SELECT ct.id, ct.number, c.name, ct.status
            FROM containers ct
            LEFT JOIN inventories i
                ON ct.id = i.id_container
            LEFT JOIN contractors c
                ON i.id_contractor = c.id
            WHERE ct.id <> 1
            GROUP BY ct.id
            """)
        self.tbv_containers.setModel(self.mdl_containers)

        self.mdl_containers.setHeaderData(1, Qt.Orientation.Horizontal, "NR PODAJNIKA")
        self.mdl_containers.setHeaderData(2, Qt.Orientation.Horizontal, "OSTATNI\nUŻYTKOWNIK")
        self.mdl_containers.setHeaderData(3, Qt.Orientation.Horizontal, "STATUS")

        self.tbv_containers.setColumnHidden(0, True)

        self.tbv_containers.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tbv_containers.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.tbv_containers.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

    # SWITCHING PAGES
    def evt_btn_inv_clicked(self):
        self.btn_inv.setDefault(True)
        self.lyt_stc.setCurrentIndex(0)

    def evt_btn_contr_clicked(self):
        self.btn_contr.setDefault(True)
        self.lyt_stc.setCurrentIndex(2)

    def evt_btn_tag_clicked(self):
        self.btn_tag.setDefault(True)
        self.lyt_stc.setCurrentIndex(1)

    # CRUD - INVENTORIES
    def evt_btn_add_clicked(self):
        dlg_add = DlgInventory(operation="Add")
        dlg_add.setModal(True)
        dlg_add.setWindowTitle("Nowe wypożyczenie")
        dlg_add.btn_return.setVisible(False)
        dlg_add.btn_add.setText("Dodaj")
        dlg_add.btn_modify.hide()

        dlg_add.show()
        dlg_add.exec_()

        self.populate_tbv_progress()
        self.populate_tbv_finished()
        self.populate_tbv_arbotags()
        self.populate_tbv_hammers()
        self.populate_tbv_history()
        self.populate_tbv_containers()
        self.evt_tab_inv_changed()

    def evt_btn_mod_clicked(self):
        if self.tab_inv.currentIndex() == 0:
            row = self.tbv_progress.currentIndex().row()
        else:
            row = self.tbv_finished.currentIndex().row()

        if row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do edycji.")
        else:
            dlg_mod = DlgInventory(operation="Modify")
            dlg_mod.setModal(True)
            if self.tab_inv.currentIndex() == 0:
                contractor = self.mdl_progress.record(row).value("name")
                dlg_mod.setWindowTitle("Edycja wypożyczenia sprzętu -\n{}".format(contractor))
            else:
                contractor = self.mdl_finished.record(row).value("name")
                dlg_mod.setWindowTitle("Zakończone wypożyczenie - {}".format(contractor))
            dlg_mod.btn_add.hide()
            dlg_mod.show()
            dlg_mod.exec_()

            self.populate_tbv_progress()
            self.populate_tbv_finished()
            self.populate_tbv_arbotags()
            self.populate_tbv_hammers()
            self.populate_tbv_history()
            self.populate_tbv_containers()
            self.evt_tab_inv_changed()

    def evt_btn_cancel_clicked(self):
        row = self.tbv_progress.currentIndex().row()
        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do anulowania")
        else:
            contractor = self.mdl_progress.record(row).value("name")
            hammer = self.mdl_progress.record(row).value("NB_HAMMER")
            container = self.mdl_progress.record(row).value("NB_CONTAINER")

            ans = translated_question("Anulowanie", "Potwierdzasz anulowanie wydania sprzętu przez:\n{} "
                                                    "(młotek {}, pojemnik {})".format(contractor,
                                                                                      hammer, container))

            if ans == QMessageBox.Yes:
                inv_id = self.mdl_progress.record(row).value("ID_INVENTORY")
                hammer_id = self.mdl_progress.record(row).value("ID_HAMMER")
                container_id = self.mdl_progress.record(row).value("ID_CONTAINER")
                arbotags_id = self.mdl_progress.record(row).value("id_arbotags")

                # UPDATE INVENTORIES TABLE
                query = QSqlQuery()
                query.prepare("UPDATE inventories SET status = 'ANULOWANE' WHERE id = :inv_id")
                query.bindValue(":inv_id", inv_id)
                b_ok_date = query.exec_()

                # UPDATE ARBOTAGS TABLE
                query.prepare("UPDATE arbotags SET id_contractor = :none, arbotags = :none, arbotags_returned = :none "
                              "WHERE id = :arbotags_id")

                query.bindValue(":arbotags_id", arbotags_id)
                query.bindValue(":none", None)
                b_ok_arbotags = query.exec_()

                # UPDATE HAMMERS TABLE
                query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                query.bindValue(":id", hammer_id)
                query.bindValue(":status", "ODDANY")
                query.exec_()

                # UPDATE CONTAINERS TABLE
                query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                query.bindValue(":id", container_id)
                query.bindValue(":status", "ODDANY")
                query.exec_()

                if b_ok_date:
                    self.populate_tbv_progress()
                    self.populate_tbv_finished()
                    self.populate_tbv_canceled()
                    self.evt_tab_inv_changed()
                else:
                    QMessageBox.critical(self, "Błąd bazy danych",
                                         "Database error\n\n{}".format(query.lastError().text()))

    def evt_btn_del_clicked(self):
        row = self.tbv_progress.currentIndex().row()
        if row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do usunięcia")
        else:
            inv_id = self.mdl_progress.record(row).value("ID_INVENTORY")
            contractor = self.mdl_progress.record(row).value("name")
            hammer = self.mdl_progress.record(row).value("id_hammer")
            container = self.mdl_progress.record(row).value("id_container")
            ans = translated_question("Usuwanie rekordu - {}".format(contractor),
                                      "Czy na pewno usunąć zaznaczoną pozycję?\n({},młotek: {}, pojemnik:{})".
                                      format(contractor, hammer, container))

            if ans == QMessageBox.Yes:
                hammer_id = self.mdl_progress.record(row).value("ID_HAMMER")
                container_id = self.mdl_progress.record(row).value("ID_CONTAINER")

                query = QSqlQuery()
                query.prepare("DELETE FROM inventories WHERE id = :id")
                query.bindValue(':id', inv_id)
                b_ok = query.exec()

                query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                query.bindValue(":id", hammer_id)
                query.bindValue(":status", "ODDANY")
                query.exec_()

                query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                query.bindValue(":id", container_id)
                query.bindValue(":status", "ODDANY")
                query.exec_()

                if b_ok:
                    self.populate_tbv_progress()
                else:
                    QMessageBox.information(self, "Usuwanie rekordu", "Brak możliwości usunięcia rekordu\n({})".
                                            format(query.lastError().text()))

    # RETURNING ARBOTAGS
    def evt_btn_tags_return(self):
        """FROM INVENTORIES MENU"""
        if self.tab_inv.currentIndex() == 0:
            row = self.tbv_progress.currentIndex().row()
            self.arbotags = self.mdl_progress.record(row).value("arbotags")
        else:
            row = self.tbv_finished.currentIndex().row()
            self.arbotags = self.mdl_finished.record(row).value("arbotags")

        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zwrotu numerów")
        else:
            if self.arbotags == "":
                QMessageBox.warning(self, "Odmowa!", "Brak wydanych numerów")
            else:
                if row == -1:
                    QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zmodyfikowania")
                else:
                    dlg_tags_return = DlgTagsReturn(origin="Inventory")
                    dlg_tags_return.show()
                    dlg_tags_return.exec_()

    def evt_btn_tags_cancel(self):
        """FROM INVENTORIES MENU"""
        if self.tab_inv.currentIndex() == 0:
            row = self.tbv_progress.currentIndex().row()
            self.arbotags_id = self.mdl_progress.record(row).value("id_arbotags")
            self.arbotags = self.mdl_progress.record(row).value("arbotags")
            self.arbotags_returned = self.mdl_progress.record(row).value("arbotags_returned")
            self.contractor = self.mdl_progress.record(row).value("name")
        else:
            row = self.tbv_finished.currentIndex().row()
            self.arbotags_id = self.mdl_finished.record(row).value("id_arbotags")
            self.arbotags = self.mdl_finished.record(row).value("arbotags")
            self.arbotags_returned = self.mdl_finished.record(row).value("arbotags_returned")
            self.contractor = self.mdl_finished.record(row).value("name")

        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do przywrócenia")
        else:
            if self.arbotags_returned == "":
                QMessageBox.warning(self, "Odmowa!", "Nie zwrócono żadnych numerów")
            else:
                tags_chk = retuning_tags_check(self.arbotags_returned)
                if tags_chk:
                    ans = translated_question("Anulowanie zwrotu arbotagów - {}".format(self.contractor),
                                              "Potwierdzasz anulowanie oddania numerów przez: "
                                              "{}? Wszystkie wyznaczone wcześniej przedziały "
                                              "numerów oddanych zostaną wliczone do listy numerów "
                                              "wydanych.".format(self.contractor))

                else:
                    ans = translated_question("Anulowanie zwrotu arbotagów - {}".format(self.contractor),
                                              "UWAGA - arbotagi znajdują się już w puli numerów wydanych.\n\n"
                                              "Anulowanie zwortu spowoduje powtórzenie numeracji w bazie danych "
                                              "numerów wydanych. W przypadku akceptacji sugerowane jest wykonanie "
                                              "późniejszej ręcznej korekty poprzez skorzystanie z funkcji "
                                              "\"Zwrot arbotagów\" lub \"Anulowanie zwrotu arbotagów.\"\n\n"
                                              "Potwierdzasz anulowanie oddania numerów przez: "
                                              "{}? Wszystkie wyznaczone wcześniej przedziały "
                                              "numerów oddanych zostaną wliczone do listy numerów "
                                              "wydanych.".format(self.contractor))

                if ans == QMessageBox.Yes:
                    # UPDATE ARBOTAGS TABLE
                    query = QSqlQuery()
                    query.prepare(
                        "UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn WHERE id = :arbotags_id")

                    tags = returned_tags_cancel(self.arbotags, self.arbotags_returned)

                    query.bindValue(":arbotags_id", self.arbotags_id)
                    query.bindValue(":tags", tags)
                    query.bindValue(":rtrn", None)
                    b_ok_arbotags = query.exec_()

                    if b_ok_arbotags:
                        self.populate_tbv_progress()
                        self.populate_tbv_finished()
                        self.populate_tbv_canceled()
                        self.populate_tbv_history()
                        self.populate_tbv_arbotags()
                    else:
                        QMessageBox.critical(self, "Błąd bazy danych",
                                             "Database error\n\n{}".format(query.lastError().text()))

    def evt_btn_atags_return(self):
        """FROM ARBOTAGS MENU"""
        idx = self.tbv_arbotags.currentIndex()
        # print(idx.siblingAtColumn(0).data())# data (id) in hidden column
        row = idx.row()

        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zwrotu numerów")
        else:
            arbotags_id = idx.siblingAtColumn(0).data()  # data (id) in hidden column
            if arbotags_id == "":
                QMessageBox.warning(self, "Odmowa!", "Brak wydanych numerów")
            else:
                if row == -1:
                    QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zmodyfikowania")
                else:
                    dlg_tags_return = DlgTagsReturn(origin="Arbotags")
                    dlg_tags_return.show()
                    dlg_tags_return.exec_()

    def evt_btn_atags_cancel(self):
        """FROM ARBOTAGS MENU"""
        idx = self.tbv_arbotags.currentIndex()
        row = idx.row()
        self.arbotags_id = idx.siblingAtColumn(0).data()
        self.arbotags = idx.siblingAtColumn(2).data()
        self.arbotags_returned = idx.siblingAtColumn(3).data()
        self.contractor = idx.siblingAtColumn(1).data()

        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do przywrócenia")
        else:
            if self.arbotags_returned == "":
                QMessageBox.warning(self, "Odmowa!", "Nie zwrócono żadnych numerów")
            else:
                tags_chk = retuning_tags_check(self.arbotags_returned)
                if tags_chk:
                    ans = translated_question("Anulowanie zwrotu arbotagów - {}".format(self.contractor),
                                              "Potwierdzasz anulowanie oddania numerów przez: "
                                              "{}? Wszystkie wyznaczone wcześniej przedziały "
                                              "numerów oddanych zostaną wliczone do listy numerów "
                                              "wydanych.".format(self.contractor))

                else:
                    ans = translated_question("Anulowanie zwrotu arbotagów - {}".format(self.contractor),
                                              "UWAGA - arbotagi znajdują się już w puli numerów wydanych.\n\n"
                                              "Anulowanie zwortu spowoduje powtórzenie numeracji w bazie danych "
                                              "numerów wydanych. W przypadku akceptacji sugerowane jest wykonanie "
                                              "późniejszej ręcznej korekty poprzez skorzystanie z funkcji "
                                              "\"Zwrot arbotagów\" lub \"Anulowanie zwrotu arbotagów.\"\n\n"
                                              "Potwierdzasz anulowanie oddania numerów przez: "
                                              "{}? Wszystkie wyznaczone wcześniej przedziały "
                                              "numerów oddanych zostaną wliczone do listy numerów "
                                              "wydanych.".format(self.contractor))

                if ans == QMessageBox.Yes:
                    # UPDATE ARBOTAGS TABLE
                    query = QSqlQuery()
                    query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn "
                                  "WHERE id = :arbotags_id")

                    tags = returned_tags_cancel(self.arbotags, self.arbotags_returned)

                    query.bindValue(":arbotags_id", self.arbotags_id)
                    query.bindValue(":tags", tags)
                    query.bindValue(":rtrn", None)
                    b_ok_arbotags = query.exec_()

                    if b_ok_arbotags:
                        self.populate_tbv_progress()
                        self.populate_tbv_finished()
                        self.populate_tbv_canceled()
                        self.populate_tbv_history()
                        self.populate_tbv_arbotags()
                    else:
                        QMessageBox.critical(self, "Błąd bazy danych",
                                             "Database error\n\n{}".format(query.lastError().text()))

    # RETURNING & RESUMING INVENTORY
    def evt_btn_return_clicked(self):
        row = self.tbv_progress.currentIndex().row()
        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zamknięcia")
        else:
            dlg_return = DlgReturn()
            dlg_return.setModal(True)
            dlg_return.show()
            dlg_return.exec_()

            self.populate_tbv_progress()
            self.populate_tbv_finished()
            self.populate_tbv_arbotags()
            self.populate_tbv_hammers()
            self.populate_tbv_containers()
            self.evt_tab_inv_changed()

    def evt_btn_resume_clicked(self):
        row = self.tbv_finished.currentIndex().row()
        if row == -1:
            QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do przywrócenia")
        else:
            contractor = self.mdl_finished.record(row).value("name")
            hammer = self.mdl_finished.record(row).value("NB_HAMMER")
            container = self.mdl_finished.record(row).value("NB_CONTAINER")
            inv_id = self.mdl_finished.record(row).value("ID_INVENTORY")
            arbotags_id = self.mdl_finished.record(row).value("id_arbotags")
            arbotags = self.mdl_finished.record(row).value("arbotags")
            arbotags_returned = self.mdl_finished.record(row).value("arbotags_returned")
            hammer_id = self.mdl_finished.record(row).value("ID_HAMMER")
            container_id = self.mdl_finished.record(row).value("ID_CONTAINER")
            hammer_status = self.mdl_finished.record(row).value("STATUS_HAMMER")
            container_status = self.mdl_finished.record(row).value("STATUS_CONTAINER")
            note = self.mdl_finished.record(row).value("note")

            if hammer_id == 1 and container_id == 1:
                QMessageBox.warning(self, "Brak możliwości przywrócenia!", "W ramach inwentaryzacji nie "
                                                                           "wypożyczono młotka ani podajnika - celem "
                                                                           "dokonania zwrotu arbotagów użyj opcji "
                                                                           "\"Zwrot arbotagów\" lub \"Anuluj zwrot "
                                                                           "arbotagów\".""")
            else:
                if hammer_status == "WYPOŻYCZONY" or container_status == "WYPOŻYCZONY":
                    QMessageBox.warning(self, "Odrzucono!", "Brak możliwości przywrócenia - młotek lub podajnik jest "
                                                            "obecnie wypożyczony.")
                else:
                    ans = translated_question("Anuluj zwrot sprzętu - {}".format(contractor),
                                              "Potwierdzasz anulowanie zdania sprzętu przez:\n{} "
                                              "(młotek: {}, pojemnik:{})?".format(contractor, hammer, container))

                    if ans == QMessageBox.Yes:
                        # UPDATE ARBOTAGS TABLE
                        query = QSqlQuery()
                        # think about it, right now there is an option to move tags between given and returned
                        # query.prepare(
                        #     "UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn WHERE id = :arbotags_id")
                        #
                        # tags = returned_tags_cancel(arbotags, arbotags_returned)
                        #
                        # query.bindValue(":arbotags_id", arbotags_id)
                        # query.bindValue(":tags", tags)
                        # query.bindValue(":rtrn", None)
                        # b_ok_arbotags = query.exec_()

                        # UPDATE INVENTORIES TABLE
                        query.prepare("UPDATE inventories SET date_return = :none, date_hammer_return =:none,"
                                      "date_container_return = :none, status = 'W TRAKCIE' WHERE id = :id")  # note =:note

                        # note = note + "\n\n{}: Cofnięty zwrot numerów:".format(datetime.datetime.today())
                        # for i in arbotags_returned.split("\n"):
                        #     note += "\n{}".format(i)

                        query.bindValue(":id", inv_id)
                        # query.bindValue(":note", note)
                        query.bindValue(":none", None)
                        b_ok_inventories = query.exec_()

                        # UPDATE HAMMERS TABLE
                        query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                        query.bindValue(":id", hammer_id)
                        query.bindValue(":status", "WYPOŻYCZONY")
                        b_ok_hammer = query.exec_()

                        # UPDATE CONTAINERS TABLE
                        query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                        query.bindValue(":id", container_id)
                        query.bindValue(":status", "WYPOŻYCZONY")
                        b_ok_container = query.exec_()

                        if b_ok_inventories and b_ok_container:  # and b_ok_arbotags:
                            self.populate_tbv_progress()
                            self.populate_tbv_finished()
                            self.populate_tbv_arbotags()
                            self.populate_tbv_hammers()
                            self.populate_tbv_containers()
                            self.evt_tab_inv_changed()
                        else:
                            QMessageBox.critical(self, "Błąd bazy danych",
                                                 "Database error\n\n{}".format(query.lastError().text()))

    # CRUD - CONTRACTORS
    def evt_btn_ctr_add_clicked(self):
        self.editing_mode_display(mode="Off")

        dlg_add = DlgContractor()
        dlg_add.setModal(True)
        dlg_add.btn_save.setVisible(False)
        dlg_add.show()
        dlg_add.exec_()
        self.populate_lsv_contractors()

    def evt_btn_ctr_mod_clicked(self):
        if self.lsv_contractors.currentIndex() is not None:
            lsv_row = self.lsv_contractors.currentIndex().row()

            if lsv_row == -1:
                QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do edycji")
            else:
                self.editing_mode_display(mode="On")

    def evt_btn_ctr_del_clicked(self):
        self.editing_mode_display(mode="Off")

        lsv_row = self.lsv_contractors.currentIndex().row()
        contractor_id = self.mdl_contractors.record(lsv_row).value("id")
        contractor_name = self.mdl_contractors.record(lsv_row).value("name")

        if lsv_row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do usunięcia")
        else:
            query = QSqlQuery()
            query.prepare("""
            SELECT COUNT(i.id_contractor) cnt FROM contractors c
            LEFT JOIN inventories i
            ON c.id = i.id_contractor
            WHERE c.id = :id""")
            query.bindValue(':id', contractor_id)
            bok = query.exec()

            if bok:
                while query.next():
                    counter = query.value(0)
            else:
                QMessageBox.critical(self, "Bład bazy danych", "Database error\n\n{}".format(query.lastError().text()))
                self.close()

            if counter > 0:
                QMessageBox.warning(self, "Brak możliwości usunięcia!",
                                    "Wykonawca jest zapisany w historii wypożyczeń sprzętu.")
            else:
                res = translated_question("Usuwanie rekordu", "Czy na pewno chcesz usunąć:\n{} (id {})?".
                                          format(contractor_name, contractor_id))
                if res == QMessageBox.Yes:
                    query = QSqlQuery()
                    query.prepare('DELETE FROM contractors WHERE id = :id')
                    query.bindValue(':id', contractor_id)
                    print(query.result().data(0))
                    b_ok = query.exec()
                    if b_ok:
                        self.populate_lsv_contractors()
                        self.populate_tbv_progress()
                        self.populate_tbv_finished()
                        self.lyt_contractors_info_reset()
                    else:
                        QMessageBox.critical(self, "Bład bazy danych", "Database error\n\n{}".
                                             format(query.lastError().text()))

    def evt_btn_ctr_save_clicked(self):
        self.editing_mode_display(mode="Off")

        query = QSqlQuery()
        query.prepare("UPDATE contractors SET id = :id, name = :name, phone = :phn, email = :em, "
                      "note = :note WHERE id = :id")

        query.bindValue(":id", self.led_id.text())
        query.bindValue(":name", self.led_name.text())
        query.bindValue(":phn", self.led_phone.text())
        query.bindValue(":em", self.led_email.text())
        query.bindValue(":note", self.ted_note.toPlainText())
        b_ok = query.exec()
        if b_ok:
            self.populate_lsv_contractors()
            self.populate_tbv_progress()
            self.populate_tbv_finished()
            self.populate_tbv_canceled()
            self.populate_tbv_history()
        else:
            QMessageBox.warning(self, "Database Error!!!!", "Database error\n({})".format(query.lastError().text()))

    def evt_btn_ctr_cancel_clicked(self):
        self.editing_mode_display(mode="Off")

    def evt_lsv_contractors_clicked(self, idx):
        self.editing_mode_display(mode="Off")
        self.btn_contr.setDefault(True)
        index = self.mdl_contractors.index(idx.row(), 1, QtCore.QModelIndex())
        name = self.mdl_contractors.index(idx.row(), 0, QtCore.QModelIndex())
        phone = self.mdl_contractors.index(idx.row(), 2, QtCore.QModelIndex())
        email = self.mdl_contractors.index(idx.row(), 3, QtCore.QModelIndex())
        note = self.mdl_contractors.index(idx.row(), 4, QtCore.QModelIndex())

        contractor_id = self.mdl_contractors.data(index, Qt.ItemDataRole.DisplayRole)
        contractor_name = self.mdl_contractors.data(name, Qt.ItemDataRole.DisplayRole)
        contractor_phone = self.mdl_contractors.data(phone, Qt.ItemDataRole.DisplayRole)
        contractor_email = self.mdl_contractors.data(email, Qt.ItemDataRole.DisplayRole)
        contractor_note = self.mdl_contractors.data(note, Qt.ItemDataRole.DisplayRole)

        self.led_id.setText(str(contractor_id))
        self.led_name.setText(contractor_name)
        self.led_phone.setText(contractor_phone)
        self.led_email.setText(contractor_email)
        self.ted_note.setPlainText(contractor_note)

        self.populate_tbv_history()

    # CRUD - HAMMERS & CONTAINERS
    def evt_btn_add_hammer_clicked(self):
        dlg_hammer = DlgEquipment(equipment="Hammer")
        dlg_hammer.setModal(True)
        dlg_hammer.setWindowTitle("Dodawanie sprzętu - młotek")
        dlg_hammer.show()
        dlg_hammer.exec_()

    def evt_btn_del_hammer_clicked(self):
        query = QSqlQuery("""
            SELECT * FROM hammers h
            INNER JOIN inventories i
            ON h.id = i.id_hammer
            """)
        query.exec_()
        self.used_hammers = []
        while query.next():
            self.used_hammers.append(query.value(0))

        idx = self.tbv_hammers.currentIndex()
        hamm_id = idx.siblingAtColumn(0).data()  # data (id) in hidden column
        hamm_nb = idx.siblingAtColumn(1).data()
        row = idx.row()

        if hamm_id in self.used_hammers:
            QMessageBox.warning(self, "Odmowa!", "Młotek został już wykorzystany w inwentaryzacji,\n"
                                                 "brak możliwości usunięcia")
        else:
            if row == -1:
                QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zwrotu numerów")
            else:
                ans = translated_question("Usuwanie numeru", "Czy na pewno chcesz usunąć młotek {}?".format(hamm_nb))

                if ans == QMessageBox.Yes:
                    query = QSqlQuery()
                    query.prepare("DELETE FROM hammers WHERE id = :id")
                    query.bindValue(":id", hamm_id)
                    b_ok = query.exec()

                    if b_ok:
                        self.populate_tbv_hammers()
                    else:
                        QMessageBox.critical(self, "Odmowa",
                                             "Błąd bazy danych\n\n{}".format(query.lastError().text()))

    def evt_btn_add_cont_clicked(self):
        dlg_cont = DlgEquipment(equipment="Container")
        dlg_cont.setModal(True)
        dlg_cont.setWindowTitle("Dodawanie sprzętu - młotek")
        dlg_cont.show()
        dlg_cont.exec_()

    def evt_btn_del_cont_clicked(self):
        query = QSqlQuery("""
                    SELECT * FROM containers ct
                    INNER JOIN inventories i
                    ON ct.id = i.id_container
                    """)
        query.exec_()
        self.used_containers = []
        while query.next():
            self.used_containers.append(query.value(0))

        idx = self.tbv_containers.currentIndex()
        cont_id = idx.siblingAtColumn(0).data()  # data (id) in hidden column
        cont_nb = idx.siblingAtColumn(1).data()
        row = idx.row()

        if cont_id in self.used_containers:
            QMessageBox.warning(self, "Odmowa!", "Podajnik został już wykorzystany w inwentaryzacji,\n"
                                                 "brak możliwości usunięcia")
        else:
            if row == -1:
                QMessageBox.warning(self, "Zaznacz rekord!", "Wskaż pozycję do zwrotu numerów")
            else:
                ans = translated_question("Usuwanie numeru", "Czy na pewno chcesz usunąć podajnik {}?".format(cont_nb))

                if ans == QMessageBox.Yes:
                    query = QSqlQuery()
                    query.prepare("DELETE FROM containers WHERE id = :id")
                    query.bindValue(":id", cont_id)
                    b_ok = query.exec()

                    if b_ok:
                        self.populate_tbv_containers()
                    else:
                        QMessageBox.critical(self, "Odmowa",
                                             "Błąd bazy danych\n\n{}".format(query.lastError().text()))

    # DISPLAY
    def lyt_contractors_info_reset(self):
        self.led_id.setText(None)
        self.led_name.setText(None)
        self.led_phone.setText(None)
        self.led_email.setText(None)
        self.ted_note.setPlainText(None)

    def evt_tab_inv_changed(self):
        if self.tab_inv.currentIndex() == 0:
            self.btn_add.setEnabled(True)
            self.btn_mod.setEnabled(True)
            self.btn_mail.setEnabled(True)
            self.btn_return.setEnabled(True)
            self.btn_cancel.setEnabled(True)
            self.btn_tags_return.setEnabled(True)
            self.btn_tags_cancel.setEnabled(True)
            self.btn_print.setEnabled(True)
            self.btn_resume.setEnabled(False)

            num = 0
            for row in range(self.mdl_progress.rowCount()):
                date_to = self.mdl_progress.record(row).value("date_to")
                if date_to != "":
                    data_tmp = date_to.split("-")
                    date_to = QDate(int("{}".format(data_tmp[0])),
                                    int("{}".format(data_tmp[1])),
                                    int("{}".format(data_tmp[2])))
                    if date_to < datetime.date.today():
                        num += 1

            self.btn_mail.setEnabled(True) if num > 0 else self.btn_mail.setEnabled(False)

        elif self.tab_inv.currentIndex() == 1:
            self.btn_return.setEnabled(False)
            self.btn_add.setEnabled(False)
            self.btn_mod.setEnabled(False)
            self.btn_mail.setEnabled(False)
            self.btn_cancel.setEnabled(False)

            self.btn_print.setEnabled(True)
            self.btn_resume.setEnabled(True)
            self.btn_tags_return.setEnabled(True)
            self.btn_tags_cancel.setEnabled(True)

        elif self.tab_inv.currentIndex() == 2:
            self.btn_return.setEnabled(False)
            self.btn_add.setEnabled(False)
            self.btn_mod.setEnabled(False)
            self.btn_mail.setEnabled(False)
            self.btn_cancel.setEnabled(False)
            self.btn_resume.setEnabled(False)
            self.btn_tags_return.setEnabled(False)
            self.btn_tags_return.setEnabled(False)
            self.btn_tags_cancel.setEnabled(False)
            self.btn_print.setEnabled(False)

    def editing_mode_display(self, mode):
        if mode == "On":
            self.btn_ctr_save.setVisible(True)
            self.btn_ctr_cancel.setVisible(True)
            self.btn_ctr_mod.setVisible(False)

            self.led_name.setEnabled(True)
            self.led_email.setEnabled(True)
            self.led_phone.setEnabled(True)
            self.ted_note.setEnabled(True)
        else:
            self.btn_ctr_save.setVisible(False)
            self.btn_ctr_cancel.setVisible(False)
            self.btn_ctr_mod.setVisible(True)

            self.led_name.setEnabled(False)
            self.led_email.setEnabled(False)
            self.led_phone.setEnabled(False)
            self.ted_note.setEnabled(False)

    # MAIL & CONTRACT
    def evt_btn_mail_clicked(self):
        dlg_mail = DlgMail()
        dlg_mail.setModal(True)
        dlg_mail.show()
        dlg_mail.exec_()

    def evt_btn_print_clicked(self):
        self.row = self.tbv_progress.currentIndex().row()
        if self.row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż umowę do wydruku.")
        else:
            dlg_print = DlgPrint()
            dlg_print.setModal(True)
            dlg_print.show()
            dlg_print.exec()


class DlgInventory(QDialog, Ui_dlg_inventory):
    def __init__(self, operation):
        super().__init__()
        self.setupUi(self)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)
        self.setLayout(self.lyt_main)
        self.operation = operation

        # MODEL SECTION
        self.mdl_tags = QStringListModel()
        self.tags_list = []
        self.returned_list = []

        # SIGNALS SECTION
        self.cmb_contractors.currentIndexChanged.connect(self.evt_cmb_contractors_changed)
        self.cmb_hammers.currentIndexChanged.connect(self.evt_cmb_hammers_changed)
        self.cmb_containers.currentIndexChanged.connect(self.evt_cmb_containers_changed)
        self.btn_add.clicked.connect(self.evt_btn_add_clicked)
        self.btn_modify.clicked.connect(self.evt_btn_modify_clicked)
        self.btn_cancel.clicked.connect(self.evt_btn_cancel_clicked)
        self.btn_return.clicked.connect(self.evt_btn_return_clicked)

        self.btn_add_tag.clicked.connect(self.evt_add_tags)
        self.btn_del_tag.clicked.connect(self.evt_del_tags)
        self.btn_mod_tag.clicked.connect(self.evt_mod_tags)
        self.btn_clear.clicked.connect(self.evt_lsv_tags_clear)

        self.led_from.textChanged.connect(self.led_tag_validation)
        self.led_to.textChanged.connect(self.led_tag_validation)

        # WIDGETS SETUP
        self.tag_menu_display()
        self.populate_dlg_inventory()
        if self.operation == "Modify":
            self.deactivate_closed_inventories()

        # DISPLAY SECTION
        self.dte_date_from.setCalendarPopup(True)
        self.dte_date_to.setCalendarPopup(True)
        self.led_from.setValidator(QIntValidator(1, 999999))
        self.led_from.setMaxLength(6)
        self.led_to.setValidator(QIntValidator(1, 999999))
        self.led_to.setMaxLength(6)
        self.lbl_id.hide()
        self.led_id.hide()

        if self.operation == "Modify":
            if dlg_main.tab_inv.currentIndex() == 0:
                self.row = dlg_main.tbv_progress.currentIndex().row()
            else:
                self.row = dlg_main.tbv_finished.currentIndex().row()

            if dlg_main.tab_inv.currentIndex() == 0:
                self.contractor = dlg_main.mdl_progress.record(self.row).value("name")
                self.contract_nb = dlg_main.mdl_progress.record(self.row).value("contract_nb")
                self.hammer = dlg_main.mdl_progress.record(self.row).value("NB_HAMMER")
                self.container = dlg_main.mdl_progress.record(self.row).value("NB_CONTAINER")
                self.arbotags = dlg_main.mdl_progress.record(self.row).value("arbotags")
                self.arbotags_returned = dlg_main.mdl_progress.record(self.row).value("arbotags_returned")
                self.date_from = dlg_main.mdl_progress.record(self.row).value("date_from")
                self.date_to = dlg_main.mdl_progress.record(self.row).value("date_to")
                self.note = dlg_main.mdl_progress.record(self.row).value("note")
            else:
                self.contractor = dlg_main.mdl_finished.record(self.row).value("name")
                self.contract_nb = dlg_main.mdl_finished.record(self.row).value("contract_nb")
                self.hammer = dlg_main.mdl_finished.record(self.row).value("NB_HAMMER")
                self.container = dlg_main.mdl_finished.record(self.row).value("NB_CONTAINER")
                self.arbotags = dlg_main.mdl_finished.record(self.row).value("arbotags")
                self.arbotags_returned = dlg_main.mdl_finished.record(self.row).value("arbotags_returned")
                self.date_from = dlg_main.mdl_finished.record(self.row).value("date_from")
                self.date_to = dlg_main.mdl_finished.record(self.row).value("date_to")
                self.note = dlg_main.mdl_finished.record(self.row).value("note")

            # CREATE LIST OF ISSUED TAGS
            self.tags_list = tags_to_list(self.arbotags)
            self.mdl_tags.setStringList(self.tags_list)
            self.lsv_tags.setModel(self.mdl_tags)

            # CREATE LIST OF RETURNED TAGS
            self.returned_list = tags_to_list(self.arbotags_returned)

            self.led_contract.setText(self.contract_nb)
            self.cmb_contractors.setCurrentText(str(self.contractor))
            self.cmb_hammers.setCurrentText(str(self.hammer))
            self.cmb_containers.setCurrentText(str(self.container))
            self.ted_note.setPlainText(self.note)

            date_tmp = self.date_from.split("-")
            date_from = QDate(int("{}".format(date_tmp[0])),
                              int(f"{date_tmp[1]}"),
                              int("{}".format(date_tmp[2])))
            self.dte_date_from.setDate(date_from)

            if self.date_to != "":
                date_tmp = self.date_to.split("-")
                date_to = QDate(int("{}".format(date_tmp[0])),
                                int(f"{date_tmp[1]}"),
                                int("{}".format(date_tmp[2])))
                self.dte_date_to.setDate(date_to)

    def deactivate_closed_inventories(self):
        if dlg_main.tab_inv.currentIndex() == 0:
            self.cmb_contractors.setEnabled(True)
            self.cmb_hammers.setEnabled(True)
            self.cmb_containers.setEnabled(True)
            self.dte_date_from.setEnabled(True)
            self.dte_date_to.setEnabled(True)
            self.ted_note.setEnabled(True)
            # self.lsv_tags.setEnabled(True)
            self.led_contract.setEnabled(True)
            self.led_to.setEnabled(True)
            self.led_from.setEnabled(True)
            self.btn_mod_tag.setEnabled(True)
            self.btn_del_tag.setEnabled(True)
            self.btn_clear.setEnabled(True)

            self.btn_return.hide()
            self.btn_modify.show()
            self.btn_cancel.show()
        else:
            self.cmb_contractors.setEnabled(False)
            self.cmb_hammers.setEnabled(False)
            self.cmb_containers.setEnabled(False)
            self.dte_date_from.setEnabled(False)
            self.dte_date_to.setEnabled(False)
            self.ted_note.setEnabled(False)
            # self.lsv_tags.setEnabled(False)
            self.led_contract.setEnabled(False)
            self.led_to.setEnabled(False)
            self.led_from.setEnabled(False)
            self.btn_mod_tag.setEnabled(False)
            self.btn_del_tag.setEnabled(False)
            self.btn_clear.setEnabled(False)

            self.btn_return.hide()
            self.btn_cancel.hide()
            self.btn_modify.hide()

    def populate_dlg_inventory(self):
        # POPULATE CMB_CONTRACTOR, LED_PHONE, LED_EMAIL
        self.mdl_contractors = QSqlQueryModel()
        self.mdl_contractors.setQuery("SELECT * FROM contractors")
        self.cmb_contractors.setModel(self.mdl_contractors)
        self.cmb_contractors.setModelColumn(1)
        self.led_phone.setText(self.mdl_contractors.record(0).value("phone"))
        self.led_email.setText(self.mdl_contractors.record(0).value("email"))

        # POPULATE CMB_HAMMERS AND CMB_CONTAINERS
        if self.operation == "Add":
            self.mdl_hammers = QSqlQueryModel()
            query = QSqlQuery("""
                SELECT * FROM hammers h 
                WHERE  status = 'ODDANY' OR h.id = 1
                GROUP BY h.id""")
            self.mdl_hammers.setQuery(query)
            self.cmb_hammers.setModel(self.mdl_hammers)
            self.cmb_hammers.setModelColumn(1)

            self.mdl_containers = QSqlQueryModel()
            query = QSqlQuery("""
                SELECT * FROM containers ct 
                WHERE ct.status = 'ODDANY' OR ct.id = 1
                GROUP BY ct.id""")
            self.mdl_containers.setQuery(query)
            self.cmb_containers.setModel(self.mdl_containers)
            self.cmb_containers.setModelColumn(1)

        elif self.operation == "Modify":
            row = dlg_main.tbv_progress.currentIndex().row()

            # POPULATE CMB_HAMMERS AND CMB_CONTAINERS
            current_hammer = dlg_main.mdl_progress.record(row).value("ID_HAMMER")
            self.mdl_hammers = QSqlQueryModel()
            query = QSqlQuery("""
                SELECT * FROM hammers h 
                WHERE status != 'WYPOŻYCZONY' OR id = {} OR h.id = 1
                GROUP BY id""".format(current_hammer))

            self.mdl_hammers.setQuery(query)
            self.cmb_hammers.setModel(self.mdl_hammers)
            self.cmb_hammers.setModelColumn(1)

            current_container = dlg_main.mdl_progress.record(row).value("ID_CONTAINER")
            self.mdl_containers = QSqlQueryModel()
            query = QSqlQuery("""
                SELECT * FROM containers ct 
                WHERE status != 'WYPOŻYCZONY' OR id = {} OR ct.id = 1
                GROUP BY id""".format(current_container))

            self.mdl_containers.setQuery(query)
            self.cmb_containers.setModel(self.mdl_containers)
            self.cmb_containers.setModelColumn(1)

        # POPULATE DTE_DATE_FROM AND DTE_DATE_TO
        today = datetime.date.today()
        self.dte_date_from.setDate(today)
        self.dte_date_to.setDate(today + datetime.timedelta(days=30))

    # DISPLAY
    def evt_cmb_contractors_changed(self, idx):
        self.contractor_id = self.mdl_contractors.data(self.mdl_contractors.index(idx, 0))
        contractor_name = self.mdl_contractors.data(self.mdl_contractors.index(idx, 1))
        contractor_phone = self.mdl_contractors.data(self.mdl_contractors.index(idx, 2))
        contractor_email = self.mdl_contractors.data(self.mdl_contractors.index(idx, 3))

        self.led_id.setText(str(self.contractor_id))
        self.led_phone.setText(contractor_phone)
        self.led_email.setText(contractor_email)

    def evt_cmb_hammers_changed(self, idx):
        self.hammer_id = self.mdl_hammers.data(self.mdl_hammers.index(idx, 0))

    def evt_cmb_containers_changed(self, idx):
        self.container_id = self.mdl_containers.data(self.mdl_containers.index(idx, 0))

    # CRUD
    def evt_btn_add_clicked(self):
        if self.hammer_id == 1 and self.container_id == 1 and self.tags_list == []:
            QMessageBox.critical(self, "Błąd!", "Nie wybrano żadnych obiektów")
        else:
            query = QSqlQuery()

            if self.hammer_id == 1 and self.container_id == 1:

                ans = translated_question("Uwaga", "Nie wybrano żadnego młotka anie podajnika. Arbotagi ze wskazanego "
                                                   "zakresu zostaną wstawione bezpośrednio do tablicy z zamkniętymi "
                                                   "inwentaryzacjami.\nCzy chcesz kontyuować?")

                if ans == QMessageBox.Yes:
                    # INSERT INTO ARBOTAGS TABLE
                    query.prepare("INSERT INTO arbotags (id_contractor, arbotags) VALUES "
                                  "(:id_cont, :tags)")
                    query.bindValue(":id_cont", self.contractor_id)

                    tags = tags_from_list(self.tags_list)
                    query.bindValue(":tags", tags)
                    b_ok_arbotags = query.exec_()

                    arbotags_id = query.lastInsertId()

                    # INSERT INTO INVENTORIES TABLE
                    today = datetime.datetime.now().date().strftime("%Y-%m-%d")
                    query.prepare("""
                        INSERT INTO inventories (contract_nb, id_contractor, id_hammer, id_container, id_arbotags, date_from, 
                        date_to, date_return, note, status) 
                        VALUES (:ctr_nb, :id_ctr, :id_hammer, :id_container,:id_arbotags, :date_from, :date_to, :date_return, 
                        :note, :status)""")

                    query.bindValue(":ctr_nb", self.led_contract.text())
                    query.bindValue(":id_ctr", int(self.contractor_id))
                    query.bindValue(":id_hammer", int(self.hammer_id))
                    query.bindValue(":id_container", int(self.container_id))
                    query.bindValue(":id_arbotags", arbotags_id)
                    query.bindValue(":date_from", self.dte_date_from.date())
                    query.bindValue(":date_to", None)
                    query.bindValue(":date_return", today)
                    query.bindValue(":note", self.ted_note.toPlainText())
                    query.bindValue(":status", "ZAKOŃCZONE")
                    b_ok_inventory = query.exec_()

                    b_ok_hammer = True
                    b_ok_container = True
                else:
                    b_ok_arbotags = True
                    b_ok_inventory = True
                    b_ok_hammer = True
                    b_ok_container = True

            else:
                # INSERT INTO HAMMERS TABLE
                query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                query.bindValue(":id", self.hammer_id)
                query.bindValue(":status", "WYPOŻYCZONY")
                b_ok_hammer = query.exec_()

                # INSERT INTO CONTAINERS TABLE
                query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                query.bindValue(":id", self.container_id)
                query.bindValue(":status", "WYPOŻYCZONY")
                b_ok_container = query.exec_()

                # INSERT INTO ARBOTAGS TABLE
                query.prepare("INSERT INTO arbotags (id_contractor, arbotags) VALUES (:id_cont, :tags)")
                query.bindValue(":id_cont", self.contractor_id)

                tags = tags_from_list(sorted(self.tags_list, reverse=True))
                query.bindValue(":tags", tags)
                b_ok_arbotags = query.exec_()

                arbotags_id = query.lastInsertId()

                # INSERT INTO INVENTORIES TABLE
                query.prepare("""
                    INSERT INTO inventories (contract_nb, id_contractor, id_hammer, id_container, id_arbotags, 
                    date_from, date_to,note, status) 
                    VALUES (:ctr_nb, :id_ctr, :id_ham, :id_cont, :id_arbotags, :date_from, :date_to, :note, :status)""")

                query.bindValue(":ctr_nb", self.led_contract.text())
                query.bindValue(":id_ctr", int(self.contractor_id))
                query.bindValue(":id_ham", int(self.hammer_id))
                query.bindValue(":id_cont", int(self.container_id))
                query.bindValue(":id_arbotags", arbotags_id)
                query.bindValue(":date_from", self.dte_date_from.date())
                query.bindValue(":date_to", self.dte_date_to.date())
                query.bindValue(":note", self.ted_note.toPlainText())
                query.bindValue(":status", "W TRAKCIE")
                b_ok_inventory = query.exec_()

            if b_ok_inventory and b_ok_hammer and b_ok_container and b_ok_arbotags:
                self.close()
            else:
                QMessageBox.critical(self, 'Błąd!', 'Błąd bazy danych\n\n{}'.format(query.lastError().text()))

    def evt_btn_modify_clicked(self):
        if self.hammer_id == 1 and self.container_id == 1 and self.tags_list == []:
            QMessageBox.critical(self, "Błąd!", "Nie wybrano żadnych obiektów")
        else:
            row = dlg_main.tbv_progress.currentIndex().row()
            inv_id = dlg_main.mdl_progress.record(row).value("ID_INVENTORY")
            contractor_id = dlg_main.mdl_progress.record(row).value("id_contractor")
            hammer_id_before = dlg_main.mdl_progress.record(row).value("ID_HAMMER")
            container_id_before = dlg_main.mdl_progress.record(row).value("ID_CONTAINER")

            query = QSqlQuery()

            if self.hammer_id == 1 and self.container_id == 1:
                ans = translated_question("Uwaga", "Nie wybrano żadnego młotka anie podajnika. Arbotagi ze wskazanego "
                                                   "zakresu zostaną wstawione bezpośrednio do tablicy z zamkniętymi "
                                                   "inwentaryzacjami.\nCzy chcesz kontyuować?")

                if ans == QMessageBox.Yes:
                    # UPDATE ARBOTAGS TABLE
                    query.prepare("UPDATE arbotags SET id_contractor = :contr , arbotags = :tags "
                                  "WHERE id = :arbotags_id")

                    tags = tags_from_list(self.tags_list)
                    tags = sorted(tags, reverse=True)
                    row = dlg_main.tbv_progress.currentIndex().row()
                    arbotags_id = dlg_main.mdl_progress.record(row).value("id_arbotags")

                    query.bindValue(":contr", self.contractor_id)
                    query.bindValue(":arbotags_id", arbotags_id)
                    query.bindValue(":tags", tags)
                    query.exec_()

                    # UPDATE INVENTORIES TABLE
                    today = datetime.datetime.now().date().strftime("%Y-%m-%d")
                    query.prepare("""
                        UPDATE inventories 
                        SET contract_nb = :nb_ctr,
                        id_contractor = :id_ctr, id_hammer = :id_hammer, id_container = :id_container, 
                        date_from = :date_from, date_to =:date_to, 
                        date_return =:date_return, note = :note, status = "ZAKOŃCZONE"
                        WHERE id = :id""")

                    query.bindValue(":id", inv_id)
                    query.bindValue(":nb_ctr", self.led_contract.text())
                    query.bindValue(":id_ctr", int(self.contractor_id))
                    query.bindValue(":id_hammer", int(self.hammer_id))
                    query.bindValue(":id_container", int(self.container_id))
                    query.bindValue(":date_from", self.dte_date_from.date())
                    query.bindValue(":date_to", None)
                    query.bindValue(":date_return", today)
                    query.bindValue(":note", self.ted_note.toPlainText())
                    b_ok = query.exec_()

                    # UPDATE HAMMERS TABLE
                    query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                    query.bindValue(":id", hammer_id_before)
                    query.bindValue(":status", "ODDANY")
                    query.exec_()

                    # UPDATE CONTAINERS TABLE
                    query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                    query.bindValue(":id", container_id_before)
                    query.bindValue(":status", "ODDANY")
                    query.exec_()

                else:
                    b_ok = True
            else:
                # UPDATE INVENTORIES TABLE
                query.prepare("""
                    UPDATE inventories 
                    SET contract_nb = :nb_ctr,
                    id_contractor = :id_ctr, id_hammer = :id_ham, id_container =:id_cont, date_from = :date_from, 
                    date_to = :date_to, note = :note WHERE id = :id""")

                query.bindValue(":id", inv_id)
                query.bindValue(":nb_ctr", self.led_contract.text())
                query.bindValue(":id_ctr", int(self.contractor_id))
                query.bindValue(":id_ham", int(self.hammer_id))
                query.bindValue(":id_cont", int(self.container_id))
                query.bindValue(":date_from", self.dte_date_from.date())
                query.bindValue(":date_to", self.dte_date_to.date())
                query.bindValue(":note", self.ted_note.toPlainText())
                b_ok = query.exec_()

                # UPDATE HAMMERS TABLE
                query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                query.bindValue(":id", hammer_id_before)
                query.bindValue(":status", "ODDANY")
                query.exec_()

                query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                query.bindValue(":id", int(self.hammer_id))
                query.bindValue(":status", "WYPOŻYCZONY")
                query.exec_()

                # UPDATE CONTAINERS TABLE
                query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                query.bindValue(":id", container_id_before)
                query.bindValue(":status", "ODDANY")
                query.exec_()

                query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                query.bindValue(":id", int(self.container_id))
                query.bindValue(":status", "WYPOŻYCZONY")
                query.exec_()

                # UPDATE ARBOTAGS TABLE
                query.prepare("UPDATE arbotags SET id_contractor = :id_contr, arbotags = :tags WHERE id = :arbotags_id")

                tags = tags_from_list(sorted(self.tags_list, reverse=True))
                row = dlg_main.tbv_progress.currentIndex().row()
                arbotags_id = dlg_main.mdl_progress.record(row).value("id_arbotags")

                query.bindValue(":id_contr", self.contractor_id)
                query.bindValue(":arbotags_id", arbotags_id)
                query.bindValue(":tags", tags)
                query.exec_()

            if b_ok:
                self.close()
            else:
                QMessageBox.critical(self, "Błąd bazy danych", "Database error\n\n{}".format(query.lastError().text()))

    def evt_btn_cancel_clicked(self):
        self.close()

    def evt_btn_return_clicked(self):
        dlg_main.evt_btn_return_clicked()
        self.close()

    def tag_menu_display(self):
        if len(self.tags_list) > 0:
            self.btn_mod_tag.setEnabled(True)
            self.btn_del_tag.setEnabled(True)
            self.btn_clear.setEnabled(True)
        else:
            self.btn_mod_tag.setEnabled(False)
            self.btn_del_tag.setEnabled(False)
            self.btn_clear.setEnabled(False)

    # ADDING ARBOTAGS
    def evt_add_tags(self):
        tags = "{}-{}".format(self.led_from.text(), self.led_to.text())
        self.tags_list.append(tags)
        self.tags_list = tags_summarizing(self.tags_list)
        self.mdl_tags.setStringList(self.tags_list)

        self.lsv_tags.setModel(self.mdl_tags)
        self.led_from.clear()
        self.led_to.clear()

        self.led_tag_validation()
        self.lbl_error.setText("")
        self.tag_menu_display()

    def evt_mod_tags(self):
        lsv_row = self.lsv_tags.currentIndex().row()
        if lsv_row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do edycji")
        else:
            tags_from = self.tags_list[lsv_row].split("-")[0]
            tags_to = self.tags_list[lsv_row].split("-")[1]
            self.led_from.setText(tags_from)
            self.led_to.setText(tags_to)

            self.tags_list.pop(lsv_row)
            self.mdl_tags.setStringList(self.tags_list)
            self.lsv_tags.setModel(self.mdl_tags)

            self.led_from.setStyleSheet("")
            self.led_to.setStyleSheet("")
            self.lbl_error.setText("")
            self.btn_add_tag.setEnabled(True)

            self.tag_menu_display()

    def evt_del_tags(self):
        lsv_row = self.lsv_tags.currentIndex().row()
        if lsv_row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do usunięcia")
        else:
            self.tags_list.pop(lsv_row)
            self.mdl_tags.setStringList(self.tags_list)
            self.lsv_tags.setModel(self.mdl_tags)

            self.tag_menu_display()

    def evt_lsv_tags_clear(self):
        self.tags_list.clear()
        self.mdl_tags.setStringList(self.tags_list)

        self.tag_menu_display()

    # VALIDATION
    def led_tag_validation(self):
        lf_chk = tags_length_check(self.led_from, self.lbl_error)
        lt_chk = tags_length_check(self.led_to, self.lbl_error)
        d_chk = tags_difference_check(self.led_from, self.led_to, self.lbl_error)
        rf_chk = tags_repeat_check(self.led_from, self.tags_list, self.lbl_error)
        rt_chk = tags_repeat_check(self.led_to, self.tags_list, self.lbl_error)
        cvr_chk = tags_cover_check(self.led_from, self.led_to, self.tags_list, self.lbl_error)

        # ALL TAGS CHECK
        all_tags = all_tags_list()
        raf_chk = tags_repeat_check(self.led_from, all_tags, self.lbl_error)
        rat_chk = tags_repeat_check(self.led_to, all_tags, self.lbl_error)
        cvra_chk = tags_cover_check(self.led_from, self.led_to, all_tags, self.lbl_error)

        if self.operation == "Modify":
            rtn_chk = returned_tags_comparison(self.led_from, self.led_to, self.returned_list, self.lbl_error)
            rtn_cvr_chk = returned_tags_cover_check(self.led_from, self.led_to, self.returned_list, self.lbl_error)
        else:
            rtn_chk = True
            rtn_cvr_chk = True

        if lf_chk and lt_chk and d_chk and rf_chk and rt_chk and cvr_chk and rtn_chk and rtn_cvr_chk and raf_chk \
                and rat_chk and cvra_chk:
            self.btn_add_tag.setEnabled(True)
            self.lbl_error.setText("")
        else:
            self.btn_add_tag.setEnabled(False)


class DlgReturn(QDialog, Ui_dlg_return):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setModal(True)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)

        self.row = dlg_main.tbv_progress.currentIndex().row()
        self.inv_id = dlg_main.mdl_progress.record(self.row).value("ID_INVENTORY")
        self.arbotags_id = dlg_main.mdl_progress.record(self.row).value("id_arbotags")
        self.arbotags = dlg_main.mdl_progress.record(self.row).value("arbotags")
        self.arbotags_returned = dlg_main.mdl_progress.record(self.row).value("arbotags_returned")
        self.hammer_id = dlg_main.mdl_progress.record(self.row).value("ID_HAMMER")
        self.container_id = dlg_main.mdl_progress.record(self.row).value("ID_CONTAINER")
        self.contractor = dlg_main.mdl_progress.record(self.row).value("name")
        self.hammer = dlg_main.mdl_progress.record(self.row).value("NB_HAMMER")
        self.container = dlg_main.mdl_progress.record(self.row).value("NB_CONTAINER")

        self.setWindowTitle("Zwrot sprzętu - {}".format(self.contractor))
        self.lbl_hammer.setText("Młotek\n({})".format(self.hammer))
        self.lbl_container.setText("Podajnik\n({})".format(self.container))

        # CREATE LIST OF TAGS GIVEN TO CONTRACTOR
        self.issued_list = tags_to_list(self.arbotags)

        # CREATE LIST OF RETURNED TAGS & MODEL SECTION
        self.returned_list = tags_to_list(self.arbotags_returned)
        self.mdl_returned = QStringListModel()

        # DISPLAY SECTION
        self.led_from.setValidator(QIntValidator(1, 999999))
        self.led_from.setMaxLength(6)
        self.led_to.setValidator(QIntValidator(1, 999999))
        self.led_to.setMaxLength(6)
        self.display_setup()
        self.checkboxes_check()

        # SIGNALS SECTION
        self.chk_equipment.toggled.connect(self.chk_equipment_toggled)
        self.chk_tags.toggled.connect(self.chk_tags_toggled)

        self.btn_add_tag.clicked.connect(self.evt_add_tags)
        # self.lsv_returned.doubleClicked.connect(self.evt_mod_tags)
        self.btn_clear.clicked.connect(self.evt_lsv_tags_clear)
        self.btn_cancel.clicked.connect(self.evt_btn_cancel_clicked)
        self.btn_save.clicked.connect(self.evt_btn_save_clicked)

        # DISABLED SECTION
        self.btn_del_tag.setHidden(True)
        self.btn_mod_tag.setHidden(True)
        # self.btn_del_tag.clicked.connect(self.evt_del_tags)
        # self.btn_mod_tag.clicked.connect(self.evt_mod_tags)

        self.led_from.textChanged.connect(self.led_tag_validation)
        self.led_to.textChanged.connect(self.led_tag_validation)

    def display_setup(self):
        if self.hammer_id == 1 and self.container_id == 1:
            self.chk_equipment.setEnabled(False)
            self.lbl_hammer.setEnabled(False)
            self.cmb_hammer.setEnabled(False)
            self.lbl_container.setEnabled(False)
            self.cmb_container.setEnabled(False)

            self.chk_equipment.setChecked(True)
            self.chk_tags.setChecked(True)
            self.led_from.setEnabled(True)
            self.led_to.setEnabled(True)
            self.btn_add_tag.setEnabled(True)
            self.btn_clear.setEnabled(True)
        elif self.hammer_id == 1 and self.container_id != 1:
            self.lbl_hammer.setEnabled(False)
            self.cmb_hammer.setEnabled(False)
        elif self.hammer_id != 1 and self.container_id == 1:
            self.lbl_container.setEnabled(False)
            self.cmb_container.setEnabled(False)

        if self.arbotags == "":
            self.btn_add_tag.setEnabled(False)
            self.chk_tags.setChecked(False)
            self.chk_tags.setEnabled(False)
            self.led_from.setEnabled(False)
            self.led_to.setEnabled(False)

    def checkboxes_check(self):
        if self.chk_equipment.isChecked() and self.chk_tags.isChecked():
            if self.returned_list == []:
                self.btn_save.setEnabled(False)
            else:
                self.btn_save.setEnabled(True)
        elif self.chk_equipment.isChecked() and not self.chk_tags.isChecked():
            self.btn_save.setEnabled(True)
        elif not self.chk_equipment.isChecked():
            self.btn_save.setEnabled(False)

    def chk_equipment_toggled(self):
        if self.chk_equipment.isChecked():
            if self.hammer_id == 1 and self.container_id == 1:
                self.lbl_hammer.setEnabled(False)
                self.cmb_hammer.setEnabled(False)
                self.lbl_container.setEnabled(False)
                self.cmb_container.setEnabled(False)
            elif self.hammer_id == 1 and self.container_id != 1:
                self.lbl_container.setEnabled(True)
                self.cmb_container.setEnabled(True)
                self.lbl_hammer.setEnabled(False)
                self.cmb_hammer.setEnabled(False)
            elif self.hammer_id != 1 and self.container_id == 1:
                self.lbl_hammer.setEnabled(True)
                self.cmb_hammer.setEnabled(True)
                self.lbl_container.setEnabled(False)
                self.cmb_container.setEnabled(False)
            else:
                self.lbl_hammer.setEnabled(True)
                self.cmb_hammer.setEnabled(True)
                self.lbl_container.setEnabled(True)
                self.cmb_container.setEnabled(True)
        else:
            self.lbl_hammer.setEnabled(False)
            self.cmb_hammer.setEnabled(False)
            self.lbl_container.setEnabled(False)
            self.cmb_container.setEnabled(False)
        self.checkboxes_check()

    def chk_tags_toggled(self):
        if self.chk_tags.isChecked():
            self.led_from.setEnabled(True)
            self.led_to.setEnabled(True)
            # self.btn_add_tag.setEnabled(True)
            self.btn_mod_tag.setEnabled(True)
            self.btn_del_tag.setEnabled(True)
            self.btn_clear.setEnabled(True)
        else:
            self.led_from.setEnabled(False)
            self.led_to.setEnabled(False)
            # self.btn_add_tag.setEnabled(False)
            self.btn_mod_tag.setEnabled(False)
            self.btn_del_tag.setEnabled(False)
            self.btn_clear.setEnabled(False)
        self.checkboxes_check()

    def evt_add_tags(self):
        tags = "{}-{}".format(self.led_from.text(), self.led_to.text())

        self.returned_list.append(tags)
        self.returned_list = tags_summarizing(self.returned_list)
        self.mdl_returned.setStringList(self.returned_list)

        self.lsv_returned.setModel(self.mdl_returned)
        self.led_from.clear()
        self.led_to.clear()

        self.led_tag_validation()
        self.lbl_error.setText("")
        self.tag_menu_display()
        self.checkboxes_check()

    def evt_mod_tags(self):
        """DISABLED FOR NOW"""
        lsv_row = self.lsv_returned.currentIndex().row()
        if lsv_row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do edycji")
        else:
            tags_from = self.returned_list[lsv_row].split("-")[0]
            tags_to = self.returned_list[lsv_row].split("-")[1]
            self.led_from.setText(tags_from)
            self.led_to.setText(tags_to)

            self.returned_list.pop(lsv_row)
            self.mdl_returned.setStringList(self.returned_list)
            self.lsv_returned.setModel(self.mdl_returned)

            self.led_from.setStyleSheet("")
            self.led_to.setStyleSheet("")
            self.btn_add_tag.setEnabled(True)

            self.tag_menu_display()

    def evt_del_tags(self):
        """DISABLED FOR NOW"""
        lsv_row = self.lsv_returned.currentIndex().row()
        if lsv_row == -1:
            QMessageBox.warning(self, "Wybierz rekord!", "Wskaż element do usunięcia")
        else:
            self.returned_list.pop(lsv_row)
            self.mdl_returned.setStringList(self.returned_list)
            self.lsv_returned.setModel(self.mdl_returned)

            self.tag_menu_display()

    def evt_lsv_tags_clear(self):
        lsv_row = self.lsv_returned.currentIndex().row()
        if lsv_row == -1:
            QMessageBox.warning(self, "Brak elementów do usunięcia!", "Lista arbotagów jest pusta")
        else:
            self.returned_list.clear()
            self.mdl_returned.setStringList(self.returned_list)
            self.btn_save.setEnabled(False)

            self.tag_menu_display()

    def evt_btn_cancel_clicked(self):
        self.close()

    def evt_btn_save_clicked(self):
        """PART OF THIS METHOD IS DISABLED FOR NOW - PREVIOUSLY IT WAS ALLOWED TO
         RETURN SINGLE EQUIPMENT AND SETTING STATUS OF 2ND AS 'EXTENDED'"""

        ans = translated_question("Zwrot sprzętu", "Potwierdzasz zdanie sprzętu przez:\n{} "
                                                   "(młotek {}, pojemnik {})".format(self.contractor,
                                                                                     self.hammer, self.container))

        if ans == QMessageBox.Yes:
            query = QSqlQuery()
            today = datetime.datetime.now().date().strftime("%Y-%m-%d")
            status_dict = {"ZWROT": "ODDANY",
                           "PRZEDŁUŻENIE": "WYPOŻYCZONY",
                           "ZNISZCZENIE/ZGUBIENIE": "ZNISZCZONY/ZGUBIONY"}

            if self.chk_equipment.isChecked():
                # if self.cmb_hammer.currentText() == "ZWROT" and self.cmb_container.currentText() == "ZWROT":
                # UPDATE INVENTORIES TABLE
                query.prepare("""   UPDATE inventories SET date_return = :today, date_hammer_return =:today,
                                    date_container_return = :today, status = 'ZAKOŃCZONE'
                                    WHERE id = :inv_id""")

                query.bindValue(":inv_id", self.inv_id)
                query.bindValue(":today", today)
                b_ok_inventory = query.exec_()

                # UPDATE HAMMERS TABLE
                query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                query.bindValue(":id", self.hammer_id)
                query.bindValue(":status", status_dict[self.cmb_hammer.currentText()])
                b_ok_hammer = query.exec_()

                # UPDATE CONTAINERS TABLE
                query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                query.bindValue(":id", self.container_id)
                query.bindValue(":status", status_dict[self.cmb_container.currentText()])
                b_ok_container = query.exec_()

                if self.chk_tags.isChecked():
                    # UPDATE ARBOTAGS TABLE
                    query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn "
                                  "WHERE id = :arbotags_id")

                    returned_tags = tags_from_list(self.returned_list)
                    if len(returned_tags) > 0:
                        tags = returned_tags.split("\n")
                        self.tags_list = tags_after_return(self.issued_list, tags)
                        tags = tags_from_list(self.tags_list)
                    else:
                        tags = tags_from_list(self.issued_list)

                    query.bindValue(":arbotags_id", self.arbotags_id)
                    query.bindValue(":tags", tags)
                    query.bindValue(":rtrn", returned_tags)
                    b_ok_arbotags = query.exec_()
                else:
                    b_ok_arbotags = True

                if b_ok_hammer and b_ok_container and b_ok_inventory and b_ok_arbotags:
                    self.close()
                else:
                    QMessageBox.critical(self, "Błąd bazy danych",
                                         "Database error\n\n{}".format(query.lastError().text()))

                # elif self.cmb_hammer.currentText() == "ZWROT" and self.cmb_container.currentText() != "ZWROT":
                #     # UPDATE INVENTORIES TABLE
                #     query.prepare("UPDATE inventories SET date_hammer_return = :today WHERE id = :inv_id")
                #     query.bindValue(":inv_id", self.inv_id)
                #     b_ok_inventory = query.exec_()
                #
                #     # UPDATE ARBOTAGS TABLE
                #     query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn "
                #                   "WHERE id = :arbotags_id")
                #
                #     returned_tags = tags_from_list(self.returned_list)
                #     if len(returned_tags) > 0:
                #         tags = returned_tags.split(";")
                #         self.tags_list = tags_after_return(self.issued_list, tags)
                #         tags = tags_from_list(self.tags_list)
                #     else:
                #         tags = tags_from_list(self.issued_list)
                #
                #     query.bindValue(":arbotags_id", self.arbotags_id)
                #     query.bindValue(":tags", tags)
                #     query.bindValue(":rtrn", returned_tags)
                #     b_ok_arbotags = query.exec_()
                #
                #     # UPDATE HAMMERS TABLE
                #     query.prepare("UPDATE hammers SET status = :status WHERE id = :id AND id != 1")
                #     query.bindValue(":id", self.hammer_id)
                #     query.bindValue(":status", status_dict[self.cmb_hammer.currentText()])
                #     b_ok_hammer = query.exec_()
                #
                #     b_ok_container = True
                #
                # elif self.cmb_hammer.currentText() != "ZWROT" and self.cmb_container.currentText() == "ZWROT":
                #     # UPDATE INVENTORIES TABLE
                #     query.prepare("UPDATE inventories SET date_container_return = :today WHERE id = :inv_id")
                #     query.bindValue(":inv_id", self.inv_id)
                #     b_ok_inventory = query.exec_()
                #
                #     # UPDATE ARBOTAGS TABLE
                #     query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn WHERE id = :arbotags_id")
                #
                #     returned_tags = tags_from_list(self.returned_list)
                #     if len(returned_tags) > 0:
                #         tags = returned_tags.split(";")
                #         self.tags_list = self.tags_after_return(self.issued_list, tags)
                #         tags = tags_from_list(self.tags_list)
                #     else:
                #         tags = tags_from_list(self.issued_list)
                #
                #     query.bindValue(":arbotags_id", self.arbotags_id)
                #     query.bindValue(":tags", tags)
                #     query.bindValue(":rtrn", returned_tags)
                #     b_ok_arbotags = query.exec_()
                #
                #     # UPDATE CONTAINERS TABLE
                #     query.prepare("UPDATE containers SET status = :status WHERE id = :id and id != 1")
                #     query.bindValue(":id", self.container_id)
                #     query.bindValue(":status", status_dict[self.cmb_container.currentText()])
                #     b_ok_container = query.exec_()
                #
                #     b_ok_hammer = True
                #
                # else:
                #     # UPDATE ARBOTAGS TABLE
                #     query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn WHERE id = :arbotags_id")
                #
                #     returned_tags = tags_from_list(returned_list)
                #     if len(returned_tags) > 0:
                #         tags = returned_tags.split(";")
                #         self.tags_list = self.tags_after_return(self.issued_list, tags)
                #         tags = tags_from_list(self.tags_list)
                #     else:
                #         tags = tags_from_list(self.issued_list)
                #
                #     query.bindValue(":arbotags_id", self.arbotags_id)
                #     query.bindValue(":tags", tags)
                #     query.bindValue(":rtrn", returned_tags)
                #     b_ok_arbotags = query.exec_()
                #
                #     b_ok_inventory = True
                #     b_ok_hammer = True
                #     b_ok_container = True

            # elif not self.chk_equipment.isChecked():
            #     # UPDATE ARBOTAGS TABLE
            #     query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn "
            #                   "WHERE id = :arbotags_id")
            #
            #     returned_tags = tags_from_list(self.returned_list)
            #     if len(returned_tags) > 0:
            #         tags = returned_tags.split(";")
            #         self.tags_list = tags_after_return(self.issued_list, tags)
            #         tags = tags_from_list(self.tags_list)
            #     else:
            #         tags = tags_from_list(self.issued_list)
            #
            #     query.bindValue(":arbotags_id", self.arbotags_id)
            #     query.bindValue(":tags", tags)
            #     query.bindValue(":rtrn", returned_tags)
            #     b_ok_arbotags = query.exec_()
            #     b_ok_inventory = True
            #
            #     if self.hammer_id == self.container_id == 1:
            #         # UPDATE INVENTORIES TABLE
            #         query.prepare("""   UPDATE inventories SET date_return = :today, status = 'ZAKOŃCZONE'
            #                                                 WHERE id = :inv_id""")
            #
            #         query.bindValue(":inv_id", self.inv_id)
            #         query.bindValue(":today", today)
            #         b_ok_inventory = query.exec_()
            #         print("inv",query.lastError().text())
            #
            #     if b_ok_arbotags and b_ok_inventory:
            #         dlg_main.populate_tbv_progress()
            #         dlg_main.populate_tbv_finished()
            #         dlg_main.evt_tab_inv_changed()
            #         dlg_main.populate_tbv_arbotags()
            #         dlg_main.populate_tbv_history()
            #         self.close()
            #     else:
            #         QMessageBox.critical(self, "Błąd bazy danych",
            #                              "Database error\n\n{}".format(query.lastError().text()))

    def tag_menu_display(self):
        if len(self.returned_list) > 0:
            self.btn_mod_tag.setEnabled(True)
            self.btn_del_tag.setEnabled(True)
            self.btn_clear.setEnabled(True)
        else:
            self.btn_mod_tag.setEnabled(False)
            self.btn_del_tag.setEnabled(False)
            self.btn_clear.setEnabled(False)

    # VALIDATION
    def led_tag_validation(self):
        lf_chk = tags_length_check(self.led_from, self.lbl_error)
        lt_chk = tags_length_check(self.led_to, self.lbl_error)
        d_chk = tags_difference_check(self.led_from, self.led_to, self.lbl_error)
        rf_chk = tags_repeat_check(self.led_from, self.returned_list, self.lbl_error)
        rt_chk = tags_repeat_check(self.led_to, self.returned_list, self.lbl_error)
        cvr_chk = tags_cover_check(self.led_from, self.led_to, self.returned_list, self.lbl_error)
        iss_chk = issued_tags_comparison(self.led_from, self.led_to, self.issued_list, self.lbl_error)
        rtn_chk = returned_tags_comparison(self.led_from, self.led_to, self.returned_list, self.lbl_error)
        rtn_cvr_chk = returned_tags_cover_check(self.led_from, self.led_to, self.returned_list, self.lbl_error)

        if lf_chk and lt_chk and d_chk and rf_chk and rt_chk and cvr_chk and iss_chk:  # and rtn_chk and rtn_cvr_chk:
            self.btn_add_tag.setEnabled(True)
            self.lbl_error.setText("")
        else:
            self.btn_add_tag.setEnabled(False)


class DlgTagsReturn(QDialog, Ui_zd_tags_return):
    def __init__(self, origin):
        super().__init__()
        self.origin = origin
        self.setupUi(self)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)
        self.setModal(True)

        if self.origin == "Inventory":
            if dlg_main.tab_inv.currentIndex() == 0:
                self.row = dlg_main.tbv_progress.currentIndex().row()
                self.arbotags_id = dlg_main.mdl_progress.record(self.row).value("id_arbotags")
                self.arbotags = dlg_main.mdl_progress.record(self.row).value("arbotags")
                self.arbotags_returned = dlg_main.mdl_progress.record(self.row).value("arbotags_returned")
                self.contractor = dlg_main.mdl_progress.record(self.row).value("name")
            else:
                self.row = dlg_main.tbv_finished.currentIndex().row()
                self.arbotags_id = dlg_main.mdl_finished.record(self.row).value("id_arbotags")
                self.arbotags = dlg_main.mdl_finished.record(self.row).value("arbotags")
                self.arbotags_returned = dlg_main.mdl_finished.record(self.row).value("arbotags_returned")
                self.contractor = dlg_main.mdl_finished.record(self.row).value("name")

        elif self.origin == "Arbotags":
            idx = dlg_main.tbv_arbotags.currentIndex()
            self.row = idx.row()
            self.arbotags_id = idx.siblingAtColumn(0).data()
            self.arbotags = idx.siblingAtColumn(2).data()
            self.arbotags_returned = idx.siblingAtColumn(3).data()
            self.contractor = idx.siblingAtColumn(1).data()

        self.setWindowTitle("Zwrot arbotagów - {}".format(self.contractor))

        # DISPLAY SECTION
        self.led_from.setValidator(QIntValidator(1, 999999))
        self.led_from.setMaxLength(6)
        self.led_to.setValidator(QIntValidator(1, 999999))
        self.led_to.setMaxLength(6)

        # CREATE LIST OF TAGS GIVEN TO CONTRACTOR
        self.issued_list = tags_to_list(self.arbotags)

        # MODEL SECTION
        self.returned_list = tags_to_list(self.arbotags_returned)

        # SIGNALS SECTION
        self.btn_add_tag.clicked.connect(self.evt_add_tags)
        self.led_from.textChanged.connect(self.led_tag_validation)
        self.led_to.textChanged.connect(self.led_tag_validation)

    def evt_add_tags(self):
        tags = "{}-{}".format(self.led_from.text(), self.led_to.text())
        self.returned_list.append(tags)
        # self.returned_list = tags_summarizing(self.returned_list)

        self.evt_btn_save()

    def evt_btn_save(self):
        ans = translated_question("Zwrot arbotagów", "Potwierdzasz oddanie arbotagów przez:\n{}".
                                  format(self.contractor))

        if ans == QMessageBox.Yes:
            # UPDATE ARBOTAGS TABLE
            query = QSqlQuery()
            query.prepare("UPDATE arbotags SET arbotags = :tags, arbotags_returned = :rtrn WHERE id = :arbotags_id")

            returned_tags = tags_from_list(self.returned_list)
            if len(returned_tags) > 0:
                tags = returned_tags.split("\n")
                self.returned_list = tags_after_return(self.issued_list, tags)
                tags = tags_from_list(self.returned_list)
            else:
                tags = tags_from_list(self.issued_list)

            # self.returned_list = tags_summarizing(self.returned_list)

            query.bindValue(":arbotags_id", self.arbotags_id)
            query.bindValue(":tags", tags)
            query.bindValue(":rtrn", returned_tags)
            b_ok_arbotags = query.exec_()

            if b_ok_arbotags:
                dlg_main.populate_tbv_progress()
                dlg_main.populate_tbv_finished()
                dlg_main.populate_tbv_canceled()
                dlg_main.populate_tbv_arbotags()
                dlg_main.populate_tbv_history()
                self.close()
            else:
                QMessageBox.critical(self, "Błąd bazy danych",
                                     "Database error\n\n{}".format(query.lastError().text()))

    # VALIDATION
    def led_tag_validation(self):
        lf_chk = tags_length_check(self.led_from, self.lbl_error)
        lt_chk = tags_length_check(self.led_to, self.lbl_error)
        d_chk = tags_difference_check(self.led_from, self.led_to, self.lbl_error)
        rf_chk = tags_repeat_check(self.led_from, self.returned_list, self.lbl_error)
        rt_chk = tags_repeat_check(self.led_to, self.returned_list, self.lbl_error)
        cvr_chk = tags_cover_check(self.led_from, self.led_to, self.returned_list, self.lbl_error)
        iss_chk = issued_tags_comparison(self.led_from, self.led_to, self.issued_list, self.lbl_error)

        if lf_chk and lt_chk and d_chk and rf_chk and rt_chk and cvr_chk and iss_chk:
            self.btn_add_tag.setEnabled(True)
            self.lbl_error.setText("")
        else:
            self.btn_add_tag.setEnabled(False)


class DlgEquipment(QDialog, Ui_dlg_equipment):
    def __init__(self, equipment):
        super().__init__()
        self.setupUi(self)
        self.equipment = equipment
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)

        if self.equipment == "Hammer":
            self.lbl_equip.setText("Nr młotka")
        else:
            self.lbl_equip.setText("Nr podajnika")

        self.led_equip.setValidator(QIntValidator(1, 99))
        self.led_equip.setMaxLength(5)
        self.lbl_error.setStyleSheet("color:red")

        self.btn_save.clicked.connect(self.evt_btn_save)
        self.btn_cancel.clicked.connect(self.evt_cancel)
        self.led_equip.textChanged.connect(self.evt_validation)

    def evt_cancel(self):
        self.close()

    def evt_btn_save(self):
        if self.equipment == "Hammer":
            query = QSqlQuery()
            query.prepare("INSERT INTO hammers (number, status )VALUES (:nb, :status)")
            query.bindValue(":nb", "nr {}".format(self.led_equip.text()))
            query.bindValue(":status", "ODDANY")
            b_ok = query.exec()
            if b_ok:
                dlg_main.populate_tbv_hammers()
                self.close()
            else:
                QMessageBox.warning(self, "Błąd!!!!", "Bład bazy danych\n({})".format(query.lastError().text()))

            self.btn_save.setEnabled(False)
        else:
            query = QSqlQuery()
            query.prepare("INSERT INTO containers (number, status )VALUES (:nb, :status)")
            query.bindValue(":nb", "nr {}".format(self.led_equip.text()))
            query.bindValue(":status", "ODDANY")
            b_ok = query.exec()
            if b_ok:
                dlg_main.populate_tbv_containers()
                self.close()
            else:
                QMessageBox.warning(self, "Błąd!!!!", "Bład bazy danych\n({})".format(query.lastError().text()))

            self.btn_save.setEnabled(False)

    def evt_validation(self):
        if self.equipment == "Hammer":
            query = QSqlQuery("SELECT * FROM hammers")
            query.exec_()

            self.hammers = []
            while query.next():
                self.hammers.append(query.value(1))

            cur_hamm = "nr {}".format(self.led_equip.text())
            if cur_hamm in self.hammers:
                self.lbl_error.setText("Wybrany młotek już istnieje")
                self.btn_save.setEnabled(False)
            else:
                self.lbl_error.setText("")
                self.btn_save.setEnabled(True)
        else:
            query = QSqlQuery("SELECT * FROM containers")
            query.exec_()

            self.containers = []
            while query.next():
                self.containers.append(query.value(1))

            cur_cont = "nr {}".format(self.led_equip.text())
            if cur_cont in self.containers:
                self.lbl_error.setText("Wybrany młotek już istnieje")
                self.btn_save.setEnabled(False)
            else:
                self.lbl_error.setText("")
                self.btn_save.setEnabled(True)


class DlgContractor(QDialog, Ui_dlg_contractor):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)

        self.btn_add.clicked.connect(self.evt_btn_add_clicked)
        self.btn_save.clicked.connect(self.evt_btn_save_clicked)
        self.btn_cancel.clicked.connect(self.evt_btn_cancel_clicked)

    def evt_btn_add_clicked(self):
        query = QSqlQuery()
        query.prepare("INSERT INTO contractors (name, phone, email, note) VALUES (:name, :phn, :em, :not)")
        query.bindValue(":name", self.led_name.text())
        query.bindValue(":phn", self.led_phone.text())
        query.bindValue(":em", self.led_email.text())
        query.bindValue(":not", self.ted_note.toPlainText())
        b_ok = query.exec()
        if b_ok:
            dlg_main.populate_tbv_progress()
            self.close()
        else:
            QMessageBox.warning(self, "Database Error!!!!", "Database error\n({})".format(query.lastError().text()))

    def evt_btn_save_clicked(self):
        query = QSqlQuery()
        query.prepare("UPDATE contractors SET id = :id, name = :name, phone = :phn, email = :em, "
                      "note = :note WHERE id = :id")

        query.bindValue(":id", self.led_id.text())
        query.bindValue(":name", self.led_name.text())
        query.bindValue(":phn", self.led_phone.text())
        query.bindValue(":em", self.led_email.text())
        query.bindValue(":note", self.ted_note.toPlainText())
        b_ok = query.exec()
        if b_ok:
            dlg_main.populate_lsv_contractors()
            dlg_main.populate_tbv_progress()
            dlg_main.populate_tbv_finished()
            dlg_main.populate_tbv_canceled()
            dlg_main.populate_tbv_history()
            self.close()
        else:
            QMessageBox.warning(self, "Database Error!!!!", "Database error\n({})".format(query.lastError().text()))

    def evt_btn_cancel_clicked(self):
        self.close()


class DlgMail(QDialog, Ui_dlg_mail):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)
        self.setWindowTitle("Wyślij przypomienie")

        self.lsv_expired.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.led_password.setEchoMode(QLineEdit.Password)

        self.led_sender.setText("123@gmail.com")
        self.led_password.setText("123")

        self.populate_lsv_expired()
        self.create_email_data()
        self.lsv_expired.selectAll()
        self.ted_mail.setPlainText(self.email_data[0][4])

        self.chk_default.toggled.connect(self.chk_default_toggled)
        self.lsv_expired.clicked.connect(self.evt_lsv_expired_clicked)
        self.ted_mail.textChanged.connect(self.evt_ted_mail_changed)
        self.btn_select_all.clicked.connect(self.evt_btn_select_all_clicked)
        self.btn_reset_changes.clicked.connect(self.evt_btn_reset_changes_clicked)
        self.btn_send.clicked.connect(self.evt_btn_send_clicked)
        self.btn_cancel.clicked.connect(self.evt_btn_cancel_clicked)

    def populate_lsv_expired(self):
        self.mdl_expired = QSqlQueryModel()
        today = datetime.datetime.now().date().strftime("%Y-%m-%d")
        query = QSqlQuery("""
            SELECT  i.id, i.contract_nb, i.id_contractor, i.id_hammer, i.date_to,
                    c.name, c.email,
                    h.number AS NB_HAMMER,
                    ct.number AS NB_CONTAINER
            FROM inventories i
            INNER JOIN contractors c
                ON i.id_contractor = c.id
            INNER JOIN hammers h
                ON i.id_hammer = h.id
            INNER JOIN containers ct
                ON i.id_container = ct.id
            WHERE i.date_to < '{}' AND NOT (h.id = 1 AND ct.id =1)
    
            INTERSECT
    
            SELECT  i.id, i.contract_nb, i.id_contractor, i.id_hammer, i.date_to,
                    c.name, c.email,
                    h.number AS NB_HAMMER,
                    ct.number AS NB_CONTAINER
            FROM inventories i
            INNER JOIN contractors c
                ON i.id_contractor = c.id
            INNER JOIN hammers h
                ON i.id_hammer = h.id
            INNER JOIN containers ct
                ON i.id_container = ct.id
            WHERE i.status = 'W TRAKCIE'""".format(today))

        query.exec()
        self.mdl_expired.setQuery(query)
        self.lsv_expired.setModel(self.mdl_expired)
        self.lsv_expired.setModelColumn(4)
        self.lsv_expired.selectAll()

    def create_email_data(self):
        self.email_data = []
        for idx in self.lsv_expired.selectedIndexes():
            self.lsv_expired.setCurrentIndex(idx)
            lsv_row = self.lsv_expired.currentIndex().row()

            inventory_id = self.mdl_expired.record(lsv_row).value("id")
            contract_nb = self.mdl_expired.record(lsv_row).value("contract_nb")
            contractor_id = self.mdl_expired.record(lsv_row).value("id_contractor")
            contractor_name = self.mdl_expired.record(lsv_row).value("name")
            contractor_email = self.mdl_expired.record(lsv_row).value("email")
            nb_hammer = self.mdl_expired.record(lsv_row).value("NB_HAMMER")
            nb_container = self.mdl_expired.record(lsv_row).value("NB_CONTAINER")
            date_to = self.mdl_expired.record(lsv_row).value("date_to")

            mail_body = (
                "Dzień dobry,\n\nPrzypominamy o konieczności zwrotu sprzętu,tj. młota: {} oraz  podajnika: {}, "
                "wypożyczonych celem przeprowadzenia inwentaryzacji drzew na terenach podległych Zarządowi XXX.\n"
                "Wyznaczony termin na oddanie ww.  sprzętu minął dnia {}.\n\nZ wyrazami szacunku,\n...").\
                format(nb_hammer, nb_container, date_to)

            mail_subject = "Zwrotu wypożyczonego sprzętu - inwentaryzacja {}".format(contract_nb)

            mail_data = [inventory_id, contractor_id, contractor_name, contractor_email, mail_body, mail_subject]
            self.email_data.append(mail_data)

    def chk_default_toggled(self, chkd):
        if chkd:
            pass
            # print("Toggled")
        else:
            pass
            # print("untoggled")

    def evt_lsv_expired_clicked(self):
        lsv_row = self.lsv_expired.currentIndex().row()
        inventory_id = self.mdl_expired.record(lsv_row).value("id")

        for mail_data in self.email_data:
            if mail_data[0] == inventory_id:
                self.ted_mail.setPlainText(mail_data[4])

    def evt_ted_mail_changed(self):
        lsv_row = self.lsv_expired.currentIndex().row()
        inventory_id = self.mdl_expired.record(lsv_row).value("id")
        for mail_data in self.email_data:
            if mail_data[0] == inventory_id:
                mail_data[4] = self.ted_mail.toPlainText()

    def evt_btn_select_all_clicked(self):
        self.lsv_expired.selectAll()
        self.evt_lsv_expired_clicked()

    def evt_btn_reset_changes_clicked(self):
        self.lsv_expired.selectAll()
        self.ted_mail.clear()
        self.create_email_data()
        self.lsv_expired.clearSelection()

    def evt_btn_cancel_clicked(self):
        self.close()

    def evt_btn_send_clicked(self):
        recipients = ""
        chk_lst = []

        if self.led_sender.text() == "" or self.led_password.text() == "":
            QMessageBox.warning(self, "Błąd", "Błędne dane nadawcy")
        else:
            for idx in self.lsv_expired.selectedIndexes():
                self.lsv_expired.setCurrentIndex(idx)

                lsv_row = self.lsv_expired.currentIndex().row()
                inventory_id = self.mdl_expired.record(lsv_row).value("id")
                contractor_id = self.mdl_expired.record(lsv_row).value("id_contractor")
                contractor_mail = self.mdl_expired.record(lsv_row).value("email")

                for mail_data in self.email_data:
                    if mail_data[0] == inventory_id:
                        # SEND MAIL HERE - TEST
                        send = send_mail(self.led_sender.text(), self.led_password.text(), self.led_sender.text(),
                                         mail_data[3], mail_data[5], mail_data[4])
                        if send:
                            chk_lst.append(True)
                        else:
                            chk_lst.append(False)

                recipients += "\n{}".format(contractor_mail)

        if False in chk_lst:
            QMessageBox.information(self, "Błąd", "Bład wysyłania wiadomości do:{}".format(recipients))
            self.lsv_expired.clearSelection()
        else:
            QMessageBox.information(self, "Wiadomość wysłana", "Wysłano przypomnienie do:{}".format(recipients))
            self.lsv_expired.clearSelection()


class DlgPrint(QDialog, Ui_dlg_print):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setWindowFlag(Qt.MSWindowsFixedSizeDialogHint)
        self.setWindowTitle("Generowanie umowy")

        # GETTING ATTRIBUTES
        self.row = dlg_main.tbv_progress.currentIndex().row()

        self.contract_nb = dlg_main.mdl_progress.record(self.row).value("contract_nb")
        self.contractor_id = dlg_main.mdl_progress.record(self.row).value("id_contractor")
        self.contractor = dlg_main.mdl_progress.record(self.row).value("contract_nb")
        self.hammer = dlg_main.mdl_progress.record(self.row).value("NB_HAMMER")
        self.container = dlg_main.mdl_progress.record(self.row).value("NB_CONTAINER")
        self.arbotags = dlg_main.mdl_progress.record(self.row).value("arbotags")
        self.date_to = dlg_main.mdl_finished.record(self.row).value("date_to")

        self.btn_save.clicked.connect(self.evt_btn_save_clicked)
        self.btn_cancel.clicked.connect(self.evt_btn_cancel_clicked)
        self.btn_template.clicked.connect(self.evt_btn_template_clicked)
        self.btn_agreement.clicked.connect(self.evt_btn_agreement_clicked)

        self.led_setup()

    def evt_btn_template_clicked(self):
        fn_in, fn_ok = QFileDialog.getOpenFileName(self, "Wczytaj plik", os.getcwd(), "Word (*.docx)")
        if fn_ok:
            self.led_template.setText(fn_in)

    def evt_btn_agreement_clicked(self):
        fn_out, fn_ok = QFileDialog.getSaveFileName(self, "Zapisz plik", os.getcwd() + r"\UMOWY", "Word (*.docx)")
        if fn_ok:
            self.led_agreement.setText(fn_out)

    def evt_btn_save_clicked(self):
        if len(self.arbotags) > 0:
            tmp_lst = [tags for tags in self.arbotags.split("\n")]
            tags = ""
            for rng in tmp_lst:
                start = int(rng.split("-")[0])
                stop = int(rng.split("-")[1])
                quantity = stop - start + 1
                tags += "{} ({} szt.); ".format(rng, quantity)
            self.tags = tags[:-2]
        else:
            self.tags = "-"

        query = QSqlQuery()
        query.prepare("SELECT phone, email FROM contractors WHERE id =:id")
        query.bindValue(":id", int(self.contractor_id))
        query.exec()

        while query.next():
            self.phone = query.value(0)
            self.email = query.value(1)

        # SETTING DIRECTORIES
        cur_path = os.getcwd()
        print_cat = r"{}\UMOWY".format(cur_path)

        if not os.path.exists(print_cat):
            os.mkdir(print_cat)

        tmpl_fn = self.led_template.text()
        agrm_fn = self.led_agreement.text()

        res = contract_from_template(tmpl_fn, agrm_fn, self.contract_nb, self.contractor, self.phone, self.email,
                                     self.hammer, self.container, self.tags, self.date_to)

        if res:
            QMessageBox.information(self, "", "Umowa została zapisana w:\n{}".format(agrm_fn))
        else:
            if not os.path.isfile(tmpl_fn):
                QMessageBox.critical(self, "Błąd", "Niewłaściwy plik szablonu.")
            elif not os.path.isdir(os.path.dirname(agrm_fn)):
                QMessageBox.critical(self, "Błąd", "Niewłaściwy katalog zapisu.")

    def evt_btn_cancel_clicked(self):
        self.close()

    def led_setup(self):
        cur_path = os.getcwd()
        print_cat = r"{}\UMOWY".format(cur_path)

        if not os.path.exists(print_cat):
            os.mkdir(print_cat)

        in_fn = cur_path + r"\contract_template.docx"
        out_fn = print_cat + r"\{}.docx".format(self.contract_nb.replace("/", "-"))

        if os.path.os.path.isfile(in_fn):
            self.led_template.setText(in_fn)
        if os.path.isdir(print_cat):
            self.led_agreement.setText(out_fn)


# SUBCLASSES OF QSqlTableModel AND QSqlQueryModel - FOR SETTING STYLE ONLY
class ProgressTableModel(QSqlTableModel):
    def __init__(self, parent=None):
        super().__init__()
        self.today = datetime.datetime.now().date().strftime("%Y-%m-%d")

    def data(self, index, role):
        if role == Qt.DisplayRole:
            return QSqlTableModel.data(self, index, Qt.DisplayRole)

        if role == Qt.BackgroundRole:
            position = self.index(index.row(), 17, QtCore.QModelIndex())
            date_to = self.data(position, Qt.ItemDataRole.DisplayRole)
            if date_to < self.today and date_to != "":
                return QBrush(QColor("red"))
            elif date_to == "":
                return QBrush(QColor("#FFD700"))

        if role == Qt.EditRole:
            return QSqlTableModel.data(self, index, Qt.DisplayRole)

        if role == Qt.ToolTipRole:
            return "Kliknij dwukrotnie żeby edytować"

        if role == Qt.TextAlignmentRole:
            return Qt.AlignCenter

    def setData(self, index, value, role=Qt.EditRole):
        return QSqlTableModel.data(self, index, Qt.DisplayRole)


class ModifiedTableModel(QSqlTableModel):
    def __init__(self, parent=None):
        super().__init__()

    def data(self, index, role):
        if role == Qt.DisplayRole:
            return QSqlTableModel.data(self, index, Qt.DisplayRole)

        if role == Qt.EditRole:
            return QSqlTableModel.data(self, index, Qt.DisplayRole)

        if role == Qt.ToolTipRole:
            return "Kliknij dwukrotnie żeby sprawdzić"

        if role == Qt.TextAlignmentRole:
            return Qt.AlignCenter

    def setData(self, index, value, role=Qt.EditRole):
        return QSqlTableModel.data(self, index, Qt.DisplayRole)


class ModifiedSqlQueryModel(QSqlQueryModel):
    def __init__(self, parent=None):
        super().__init__()

    def data(self, index, role):
        if role == Qt.DisplayRole:
            return QSqlQueryModel.data(self, index, Qt.DisplayRole)

        if role == Qt.TextAlignmentRole:
            return Qt.AlignCenter

    def setData(self, index, value, role=Qt.EditRole):
        return QSqlTableModel.data(self, index, Qt.DisplayRole)


# TAGS CALCULATION FUNCTIONS
def tags_to_list(tags_text):
    result_list = []
    tags_list = tags_text.split("\n") if len(tags_text) > 0 else []
    for tags in tags_list:
        result_list.append(tags)

    return result_list


def tags_from_list(tags_list):
    tags = ""
    if len(tags_list) > 0:
        for i in sorted(tags_list, reverse=True):
            tags += "{}\n".format(i)
        tags = tags.strip("\n")
    # else:
    #     tags = ""
    return tags


def tags_after_return(issued_list, returned_list):
    new_issued_list = issued_list.copy()

    for i in range(len(returned_list)):
        returned_range = returned_list[i].split("-")
        returned_start = int(returned_range[0])
        returned_stop = int(returned_range[1])

        for j in range(len(issued_list)):
            # issued_tags_range = new_issued_list[j].split("-")
            issued_tags_range = issued_list[j].split("-")
            issued_start = int(issued_tags_range[0])
            issued_stop = int(issued_tags_range[1])

            if len(new_issued_list) > 0:
                print(new_issued_list)

            if returned_start in range(issued_start, issued_stop + 1) and \
                    returned_stop in range(issued_start, issued_stop + 1):
                # FULL RETURN
                if returned_start == issued_start and returned_stop == issued_stop:
                    del_tags = "{}-{}".format(issued_start, issued_stop)
                    if del_tags in new_issued_list:
                        new_issued_list.remove(del_tags)

                # START THE SAME, DIFFERENT STOP
                elif returned_start == issued_start and returned_stop != issued_stop:
                    add_tags = "{}-{}".format(issued_stop - (issued_stop - returned_stop - 1), issued_stop)
                    new_issued_list.append(add_tags)

                    del_tags = "{}-{}".format(issued_start, issued_stop)
                    if del_tags in new_issued_list:
                        new_issued_list.remove(del_tags)

                # START DIFFERENT, STOP THE SAME
                elif returned_start != issued_start and returned_stop == issued_stop:
                    add_tags = "{}-{}".format(issued_start, issued_start + (returned_start - issued_start - 1))
                    new_issued_list.append(add_tags)

                    del_tags = "{}-{}".format(issued_start, issued_stop)
                    if del_tags in new_issued_list:
                        new_issued_list.remove(del_tags)

                # START DIFFERENT, STOP DIFFERENT
                elif returned_start != issued_start and returned_stop != issued_stop:
                    add_tags_1 = "{}-{}".format(issued_stop - (issued_stop - returned_stop - 1), issued_stop)
                    add_tags_2 = "{}-{}".format(issued_start, issued_start + (returned_start - issued_start - 1))
                    new_issued_list.append(add_tags_1)
                    new_issued_list.append(add_tags_2)

                    del_tags = "{}-{}".format(issued_start, issued_stop)
                    if del_tags in new_issued_list:
                        new_issued_list.remove(del_tags)

    return sorted(new_issued_list, reverse=True)


def tags_summarizing(tags_list):
    tags_list.sort()
    lst_new_tags = []
    lst_to_delete = []

    for i in range(len(tags_list)):
        if tags_list[i] != '':
            start = int(tags_list[i].split("-")[0])
            stop = int(tags_list[i].split("-")[1])

            if i < len(tags_list) - 1:
                start_next = int(tags_list[i + 1].split("-")[0])
                stop_next = int(tags_list[i + 1].split("-")[1])

                if start_next - stop == 1:
                    tags = "{}-{}".format(start, stop_next)
                    lst_new_tags.append(tags)

                    repeated_tags = "{}-{}".format(start_next, stop_next)
                    lst_to_delete.append(repeated_tags)
                else:
                    tags = "{}-{}".format(start, stop)
                    lst_new_tags.append(tags)
            else:
                tags = "{}-{}".format(start, stop)
                lst_new_tags.append(tags)

    for i in lst_to_delete:
        if i in lst_new_tags:
            lst_new_tags.remove(i)

    # lst_new_tags = list(dict.fromkeys(lst_new_tags))
    lst_new_tags.sort()

    return lst_new_tags


def returned_tags_cancel(tags, tags_returned):
    list_tmp = tags.split("\n")
    lst_tags = []
    for i in sorted(list_tmp):
        lst_tags.append(i.strip())

    lst_tags_returned = tags_returned.split("\n")
    for i in sorted(lst_tags_returned):
        lst_tags.append(i.strip())

    # EACH NUMBER ADDED HAS TO BE SUMMARIZED SEPARATELY!
    sum_tags = []
    for i in sorted(lst_tags):
        sum_tags.append(i)
        sum_tags = tags_summarizing(sum_tags)
    tags = tags_from_list(sum_tags)

    return tags


# TAGS VALIDATION FUNCTIONS
def tags_length_check(led, lbl_error):
    text = led.text()
    if len(text) < 6 and len(text) != 0:
        led.setStyleSheet(style_led_tag_error())
        lbl_error.setStyleSheet("color:red")
        lbl_error.setText("Wymagane 6 znaków.")
    elif len(text) == 6 or len(text) == 0:
        led.setStyleSheet("")
    if len(text) == 6:
        return True


def tags_difference_check(led_from, led_to, lbl_error):
    tag_from = led_from.text()
    tag_to = led_to.text()
    if len(tag_from) == 6 and len(tag_to) == 6:
        if int(tag_to) - int(tag_from) >= 0:
            led_from.setStyleSheet("")
            led_to.setStyleSheet("")
            return True
        else:
            led_from.setStyleSheet(style_led_tag_error())
            led_to.setStyleSheet(style_led_tag_error())

            lbl_error.setStyleSheet("color:red")
            lbl_error.setText("Wartość końcowa jest większa od początkowej")


def tags_repeat_check(led, tags_list, lbl_error):
    tag = led.text()
    if len(tag) == 0:
        led.setStyleSheet("")
        return True
    else:
        if len(tags_list) == 0:
            return True
        else:
            if len(tag) == 6:
                check_list = []
                for i in range(len(tags_list)):
                    tags_list[i] = tags_list[i].replace("\n", "")
                    tags_range = tags_list[i].split("-")
                    range_start = int(tags_range[0])
                    range_stop = int(tags_range[1]) + 1

                    if int(tag) in range(range_start, range_stop):
                        led.setStyleSheet(style_led_tag_error())
                        check_list.append(False)
                    else:
                        check_list.append(True)

                if False in check_list:
                    led.setStyleSheet(style_led_tag_error())
                    lbl_error.setStyleSheet("color:red")
                    lbl_error.setText("Podany przedział powtarza zwrócone/dodane\nwcześniej wartości.")
                    return False
                else:
                    return True


def tags_cover_check(led_from, led_to, tags_list, lbl_error):
    tag_from = led_from.text()
    tag_to = led_to.text()
    if len(tags_list) == 0:
        return True
    elif len(tags_list) > 0:
        if len(tag_from) and len(tag_to) == 6:
            check_list = []
            for i in range(len(tags_list)):
                tags_range = tags_list[i].split("-")
                range_start = int(tags_range[0])
                range_stop = int(tags_range[1])

                if range_start in range(int(tag_from), int(tag_to)):
                    check_list.append(False)
                elif range_stop in range(int(tag_from), int(tag_to)):
                    check_list.append(False)
                else:
                    check_list.append(True)

            if False in check_list:
                led_from.setStyleSheet(style_led_tag_error())
                led_to.setStyleSheet(style_led_tag_error())
                lbl_error.setStyleSheet("color:red")
                lbl_error.setText("Podane wartości obejmują wydany wcześniej zakres numerów.")
                return False
            else:
                return True


def returned_tags_cover_check(led_from, led_to, tags_list, lbl_error):
    """JUST TO MAKE DIFFERENCE IN LBL_ERROR BETWEEN ISSUED AND RETURNED
    TAGS IN DLGINVENTORY"""

    tag_from = led_from.text()
    tag_to = led_to.text()
    if len(tags_list) == 0:
        return True
    elif len(tags_list) > 0:
        if len(tag_from) and len(tag_to) == 6:
            check_list = []
            for i in range(len(tags_list)):
                tags_range = tags_list[i].split("-")
                range_start = int(tags_range[0])
                range_stop = int(tags_range[1])

                if range_start in range(int(tag_from), int(tag_to)):
                    check_list.append(False)
                elif range_stop in range(int(tag_from), int(tag_to)):
                    check_list.append(False)
                else:
                    check_list.append(True)

            if False in check_list:
                led_from.setStyleSheet(style_led_tag_error())
                led_to.setStyleSheet(style_led_tag_error())
                lbl_error.setStyleSheet("color:red")
                lbl_error.setText("Podane wartości obejmują zwrócony wcześniej zakres numerów.")
                return False
            else:
                return True


def issued_tags_comparison(led_from, led_to, issued_list, lbl_error):
    check_list = []

    if len(led_from.text()) < 6 or len(led_to.text()) < 6:
        return True
    elif len(led_from.text()) == 6 and len(led_to.text()) == 6:
        for i in range(len(issued_list)):
            issued_tags_range = issued_list[i].split("-")
            issued_start = int(issued_tags_range[0])
            issued_stop = int(issued_tags_range[1]) + 1

            if int(led_from.text()) in range(issued_start, issued_stop) and \
                    int(led_to.text()) in range(issued_start, issued_stop):
                check_list.append(True)

        if True in check_list:
            return True
        else:
            led_from.setStyleSheet(style_led_tag_error())
            led_to.setStyleSheet(style_led_tag_error())
            lbl_error.setStyleSheet("color:red")
            lbl_error.setText("Wartości wykraczają poza wydany zakres arbotagów.")
            return False


def returned_tags_comparison(led_from, led_to, returned_list, lbl_error):
    check_list = []

    if len(returned_list) == 0:
        return True
    else:
        if len(led_from.text()) < 6 or len(led_to.text()) < 6:
            return True
        elif len(led_from.text()) == 6 and len(led_to.text()) == 6:
            for i in range(len(returned_list)):
                returned_tags_range = returned_list[i].split("-")
                returned_start = int(returned_tags_range[0])
                returned_stop = int(returned_tags_range[1]) + 1

                if int(led_from.text()) in range(returned_start, returned_stop) or \
                        int(led_to.text()) in range(returned_start, returned_stop):
                    check_list.append(False)

            if False not in check_list:
                return True
            else:
                led_from.setStyleSheet(style_led_tag_error())
                led_to.setStyleSheet(style_led_tag_error())
                lbl_error.setStyleSheet("color:red")
                lbl_error.setText("Numery z podanego zakresu został już zwrócone.")
                return False


# RETURNING TAGS VALIDATION - added in the end, quite complicated and not optimized
def retuning_tags_check(arbotags_returned):
    def returned_repeat_check(tag, tags_list):
        if len(tag) == 0:
            return True
        else:
            if len(tags_list) == 0:
                return True
            else:
                if len(tag) == 6:
                    check_list = []
                    for i in range(len(tags_list)):
                        tags_list[i] = tags_list[i].replace("\n", "")
                        tags_range = tags_list[i].split("-")
                        range_start = int(tags_range[0])
                        range_stop = int(tags_range[1]) + 1

                        if int(tag) in range(range_start, range_stop):
                            check_list.append(False)
                        else:
                            check_list.append(True)

                    if False in check_list:
                        return False
                    else:
                        return True

    def returned_cover_check(tag_from, tag_to, tags_list):
        if len(tags_list) == 0:
            return True
        elif len(tags_list) > 0:
            if len(tag_from) and len(tag_to) == 6:
                check_list = []
                for i in range(len(tags_list)):
                    tags_range = tags_list[i].split("-")
                    range_start = int(tags_range[0])
                    range_stop = int(tags_range[1])

                    if range_start in range(int(tag_from), int(tag_to)):
                        check_list.append(False)
                    elif range_stop in range(int(tag_from), int(tag_to)):
                        check_list.append(False)
                    else:
                        check_list.append(True)

                if False in check_list:
                    return False
                else:
                    return True

    chk_lst = []
    all_tags = all_tags_list()
    rtn = tags_to_list(arbotags_returned)
    for i in rtn:
        start = i.split("-")[0]
        stop = i.split("-")[1]

        raf_chk = returned_repeat_check(start, all_tags)
        rat_chk = returned_repeat_check(stop, all_tags)
        cvra_chk = returned_cover_check(start, stop, all_tags)

        if not (raf_chk or rat_chk or cvra_chk):
            chk_lst.append(False)

    if False in chk_lst:
        return False
    else:
        return True


# ALL TAGS LIST CREATOR
def all_tags_list():
    text_tags = []
    nested_tags = []

    query = QSqlQuery()
    query.prepare("SELECT arbotags, arbotags_returned FROM arbotags")
    query.exec()

    while query.next():
        tags = query.value(0)
        text_tags.append(tags)

    for i in text_tags:
        nested_tags.append(tags_to_list(i))

    all_tags = []
    for i in nested_tags:
        for j in i:
            if j not in all_tags:
                all_tags.append(j)

    return all_tags


# EMAIL SENDING FUNCTION (GMAIL)
def send_mail(user, password, mail_from, mail_to, mail_subject, mail_body):
    message = '''From: {}
    Dotyczy: {}\n\n{}
    '''.format(mail_from, mail_subject, mail_body)


    try:
        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.ehlo()
        server.login(user, password)
        server.sendmail(user, mail_to, message.encode("utf-8"))
        server.close()
        return True
    except:
        return False


# CONTRACT CREATOR
def contract_from_template(in_fn, out_fn, contract_nb, contractor, phone, email, hammer, container, tags, date_to):
    values_map = {"%CONTRACT%": contract_nb,
                  "%CONTRACTOR%": contractor,
                  "%PHONE%": phone,
                  "%EMAIL%": email,
                  "%HAMMER%": hammer,
                  "%CONTAINER%": container,
                  "%TAGS%": tags,
                  "%DATE_TO%": date_to,
                  }
    try:
        document = Document(in_fn)
        for nb, par in enumerate(document.paragraphs):
            for k, v in values_map.items():
                if k in par.text:
                    new_text = document.paragraphs[nb].text.replace(k, v)
                    document.paragraphs[nb].text = new_text
        document.save(out_fn)
        return True
    except:
        return False


# ADDITIONAL FUNCTIONS
def translated_question(title, text):
    box = QMessageBox()
    box.setIcon(QMessageBox.Question)
    box.setWindowTitle(title)
    box.setText(text)
    box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
    btn_y = box.button(QMessageBox.Yes)
    btn_y.setText('Tak')
    btn_n = box.button(QMessageBox.No)
    btn_n.setText('Nie')
    box.setDefaultButton(QMessageBox.No)
    result = box.exec_()

    return result


def style_led_tag_error():
    style = """
        QLineEdit {
            border: 1px solid red
        }    
    """
    return style


if __name__ == "__main__":
    app = QApplication(sys.argv)
    dlg_main = DlgMain()
    dlg_main.show()
    sys.exit(app.exec_())
